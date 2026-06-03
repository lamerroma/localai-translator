import io
import json
import os
import sqlite3
import logging
import traceback
import datetime
import uuid
import threading
import queue as _queue_mod
import secrets
import base64
import time
from urllib.parse import quote
import requests as req_lib
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, Response
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s %(levelname)s %(message)s')
log = logging.getLogger("translator")

# ── Configuration ────────────────────────────────────────────────────────────

CONFIG_FILE = os.path.join(os.path.dirname(__file__), "translator_config.json")

DEFAULTS = {
    "base_url":        "http://127.0.0.1:11434/v1",
    "model":           "rinex20/translategemma3:12b",
    "max_tokens":      2048,
    "llm_timeout":     180,
    "chunk_size":      3000,
    "auth_user":       "admin",
    "auth_pass":       "translate",
    "max_pdf_pages":   10,
    "max_chars":       30000,
}

HOST = "0.0.0.0"
PORT = 7860

LANG_NAMES_UK = {
    "Arabic":     "Арабська",
    "Bulgarian":  "Болгарська",
    "Chinese":    "Китайська",
    "Czech":      "Чеська",
    "Danish":     "Данська",
    "Dutch":      "Нідерландська",
    "English":    "Англійська",
    "Esperanto":  "Есперанто",
    "Finnish":    "Фінська",
    "French":     "Французька",
    "German":     "Німецька",
    "Greek":      "Грецька",
    "Gujarati":   "Гуджараті",
    "Hebrew":     "Іврит",
    "Hindi":      "Гінді",
    "Hungarian":  "Угорська",
    "Indonesian": "Індонезійська",
    "Italian":    "Італійська",
    "Japanese":   "Японська",
    "Korean":     "Корейська",
    "Latin":      "Латинська",
    "Persian":    "Перська",
    "Polish":     "Польська",
    "Portuguese": "Португальська",
    "Romanian":   "Румунська",
    "Russian":    "Російська",
    "Slovak":     "Словацька",
    "Spanish":    "Іспанська",
    "Swedish":    "Шведська",
    "Tagalog":    "Тагальська",
    "Thai":       "Тайська",
    "Turkish":    "Турецька",
    "Ukrainian":  "Українська",
    "Vietnamese": "В'єтнамська",
}

LANG_MAP = {
    "Arabic":     "ar",
    "Bulgarian":  "bg",
    "Chinese":    "zh",
    "Czech":      "cs",
    "Danish":     "da",
    "Dutch":      "nl",
    "English":    "en",
    "Esperanto":  "eo",
    "Finnish":    "fi",
    "French":     "fr",
    "German":     "de",
    "Greek":      "el",
    "Gujarati":   "gu",
    "Hebrew":     "he",
    "Hindi":      "hi",
    "Hungarian":  "hu",
    "Indonesian": "id",
    "Italian":    "it",
    "Japanese":   "ja",
    "Korean":     "ko",
    "Latin":      "la",
    "Persian":    "fa",
    "Polish":     "pl",
    "Portuguese": "pt",
    "Romanian":   "ro",
    "Russian":    "ru",
    "Slovak":     "sk",
    "Spanish":    "es",
    "Swedish":    "sv",
    "Tagalog":    "tl",
    "Thai":       "th",
    "Turkish":    "tr",
    "Ukrainian":  "uk",
    "Vietnamese": "vi",
}



def load_config() -> dict:
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, encoding="utf-8") as f:
                data = json.load(f)
            return {**DEFAULTS, **data}
        except Exception:
            pass
    return dict(DEFAULTS)


def save_config(cfg: dict) -> None:
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


CFG = load_config()

# ── Stats DB ─────────────────────────────────────────────────────────────────

STATS_DB = os.path.join(os.path.dirname(__file__), "stats.db")
_stats_lock = threading.Lock()


def _stats_conn():
    conn = sqlite3.connect(STATS_DB)
    conn.row_factory = sqlite3.Row
    return conn


def init_stats_db():
    with _stats_conn() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS stats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT NOT NULL,
                ip TEXT,
                kind TEXT NOT NULL,
                filename TEXT,
                file_ext TEXT,
                lang_from TEXT,
                lang_to TEXT,
                chars INTEGER,
                pages INTEGER,
                duration REAL,
                status TEXT NOT NULL,
                error TEXT
            )
        """)
        conn.commit()


def log_stat(**kwargs):
    fields = ["timestamp", "ip", "kind", "filename", "file_ext",
              "lang_from", "lang_to", "chars", "pages", "duration",
              "status", "error"]
    values = [kwargs.get(f) for f in fields]
    if not values[0]:
        values[0] = datetime.datetime.now().isoformat(timespec="seconds")
    try:
        with _stats_lock, _stats_conn() as conn:
            conn.execute(
                f"INSERT INTO stats ({','.join(fields)}) "
                f"VALUES ({','.join(['?'] * len(fields))})",
                values,
            )
            conn.commit()
    except Exception as e:
        log.error(f"log_stat failed: {e}")


def get_recent_stats(limit=100):
    with _stats_conn() as conn:
        rows = conn.execute(
            "SELECT * FROM stats ORDER BY id DESC LIMIT ?", (limit,)
        ).fetchall()
    return [dict(r) for r in rows]


def get_stats_summary():
    with _stats_conn() as conn:
        row = conn.execute("""
            SELECT
              COUNT(*) AS total,
              SUM(CASE WHEN status='success' THEN 1 ELSE 0 END) AS success,
              SUM(CASE WHEN status='error' THEN 1 ELSE 0 END) AS errors,
              SUM(CASE WHEN status='stopped' THEN 1 ELSE 0 END) AS stopped,
              SUM(chars) AS total_chars,
              SUM(duration) AS total_seconds,
              SUM(pages) AS total_pages
            FROM stats
        """).fetchone()
    return dict(row) if row else {}


init_stats_db()

# ─────────────────────────────────────────────────────────────────────────────

app = FastAPI()


ADMIN_PATHS = ("/admin", "/config")


@app.middleware("http")
async def basic_auth(request: Request, call_next):
    path = request.url.path
    if not any(path.startswith(p) for p in ADMIN_PATHS):
        return await call_next(request)
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Basic "):
        try:
            decoded = base64.b64decode(auth[6:]).decode("utf-8")
            user, _, pwd = decoded.partition(":")
            if (secrets.compare_digest(user, CFG.get("auth_user", "admin")) and
                    secrets.compare_digest(pwd, CFG.get("auth_pass", "translate"))):
                return await call_next(request)
        except Exception:
            pass
    return Response(
        status_code=401,
        headers={"WWW-Authenticate": 'Basic realm="Translator Admin"'},
    )


_active: dict[str, threading.Event] = {}
_results: dict[str, tuple[str, bytes]] = {}

_job_queue: _queue_mod.Queue = _queue_mod.Queue()
_ticket_lock = threading.Lock()
_ticket_issued = 0
_ticket_serving = 0


def _queue_worker():
    global _ticket_serving
    while True:
        start_evt, done_evt = _job_queue.get()
        with _ticket_lock:
            _ticket_serving += 1
        start_evt.set()
        done_evt.wait()


threading.Thread(target=_queue_worker, daemon=True).start()


def extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        return "\n\n".join(pages)
    elif ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        return content.decode("utf-8", errors="replace")


def split_chunks(text: str) -> list[str]:
    max_chars = CFG["chunk_size"]
    paragraphs = text.split("\n\n")
    chunks, current, current_len = [], [], 0
    for para in paragraphs:
        if current and current_len + len(para) > max_chars:
            chunks.append("\n\n".join(current))
            current, current_len = [para], len(para)
        else:
            current.append(para)
            current_len += len(para)
    if current:
        chunks.append("\n\n".join(current))
    return [c for c in chunks if c.strip()]


def call_llm(messages: list, stop_event: threading.Event) -> str | None:
    resp = req_lib.post(
        f"{CFG['base_url']}/chat/completions",
        headers={"Authorization": "Bearer dummy", "Content-Type": "application/json"},
        json={"model": CFG["model"], "stream": True, "max_tokens": CFG["max_tokens"], "messages": messages},
        timeout=CFG["llm_timeout"],
        stream=True,
    )
    collected = []
    for raw_line in resp.iter_lines():
        if stop_event.is_set():
            resp.close()
            return None
        if not raw_line:
            continue
        line = raw_line.decode("utf-8") if isinstance(raw_line, bytes) else raw_line
        if not line.startswith("data: "):
            continue
        payload = line[6:]
        if payload.strip() == "[DONE]":
            break
        try:
            chunk = json.loads(payload)
        except json.JSONDecodeError:
            continue
        token = chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
        if token:
            collected.append(token)
    raw = "".join(collected).strip()
    result = raw.split("§")[0].strip()
    return result.replace("⟨P⟩", "\n\n").replace("⟨N⟩", "\n")


def _ollama_native_host() -> str:
    """Strip the /v1 suffix used by OpenAI-compatible endpoint."""
    url = CFG["base_url"].rstrip("/")
    if url.endswith("/v1"):
        url = url[:-3]
    return url


def _translate_unit(text: str, lang_from: str, lang_to: str,
                    stop_event: threading.Event) -> str | None:
    """Translate a single short piece of text (one run / one block).

    Uses native Ollama /api/generate with a minimal prompt — much faster than
    the OpenAI-compatible /v1/chat/completions path with the heavy template,
    because for short units the template was costing more tokens than the text.
    """
    if not text or len(text.strip()) < 2:
        return text
    if stop_event.is_set():
        return None
    # Use /api/chat so Ollama applies the model's baked-in chat template
    # (Open WebUI does the same; /api/generate skips the template and the
    # model was trained to expect it).
    target = lang_to if lang_to else "Ukrainian"
    # Natural-language instruction matches what works for users in Open WebUI;
    # the model card's "To <Lang>:" anchor turned out to give worse term
    # choices than a plain "Translate to <Lang>" prompt.
    content = f"Translate to {target}:\n\n{text}"
    try:
        resp = req_lib.post(
            f"{_ollama_native_host()}/api/chat",
            json={
                "model": CFG["model"],
                "messages": [{"role": "user", "content": content}],
                "stream": False,
            },
            timeout=CFG["llm_timeout"],
        )
        if stop_event.is_set():
            return None
        if resp.status_code != 200:
            log.warning(f"Ollama returned {resp.status_code}: {resp.text[:200]}")
            return text
        data = resp.json()
        return (data.get("message", {}).get("content") or "").strip() or text
    except req_lib.exceptions.Timeout:
        raise
    except Exception as e:
        log.warning(f"_translate_unit failed: {e}")
        return text


def _parse_html_to_parts(html: str):
    """Parse HTML into (parts, text_indices, batch_text) for translation."""
    from html.parser import HTMLParser

    parts = []

    class _Parser(HTMLParser):
        def handle_starttag(self, tag, attrs):
            parts.append(('tag', self.get_starttag_text()))
        def handle_endtag(self, tag):
            parts.append(('tag', f'</{tag}>'))
        def handle_data(self, data):
            if data.strip():
                parts.append(('text', data))
            else:
                parts.append(('tag', data))
        def handle_entityref(self, name):
            parts.append(('tag', f'&{name};'))
        def handle_charref(self, name):
            parts.append(('tag', f'&#{name};'))

    _Parser().feed(html)
    text_indices = [i for i, (t, _) in enumerate(parts) if t == 'text']
    batch = ''.join(f'⟦{i}⟧{parts[i][1]}' for i in text_indices)
    return parts, text_indices, batch


def _reconstruct_html(parts: list, text_indices: list, translated_batch: str) -> str:
    """Put translated text nodes back into HTML structure using ⟦N⟧ markers."""
    for idx in text_indices:
        marker = f'⟦{idx}⟧'
        start = translated_batch.find(marker)
        if start == -1:
            continue
        start += len(marker)
        next_marker = translated_batch.find('⟦', start)
        end = next_marker if next_marker != -1 else len(translated_batch)
        parts[idx] = ('text', translated_batch[start:end])
    return ''.join(content for _, content in parts)


def _translate_html_nodes(html: str, lang_from: str, lang_to: str,
                          stop_event: threading.Event) -> str:
    """Translate HTML preserving all tags. Text nodes only, single Ollama call."""
    parts, text_indices, batch = _parse_html_to_parts(html)
    if not text_indices:
        return html
    translated = _translate_unit(batch, lang_from, lang_to, stop_event)
    if not translated:
        return html
    return _reconstruct_html(parts, text_indices, translated)


def _translate_unit_streaming(text: str, lang_from: str, lang_to: str,
                              stop_event: threading.Event):
    """Generator that yields tokens from Ollama stream one by one."""
    if not text or len(text.strip()) < 2:
        yield text
        return
    if stop_event.is_set():
        return

    target = lang_to if lang_to else "Ukrainian"
    content = f"Translate to {target}:\n\n{text}"

    try:
        resp = req_lib.post(
            f"{_ollama_native_host()}/api/chat",
            json={
                "model": CFG["model"],
                "messages": [{"role": "user", "content": content}],
                "stream": True,
            },
            timeout=CFG["llm_timeout"],
            stream=True,
        )
        if resp.status_code != 200:
            log.warning(f"Ollama stream returned {resp.status_code}")
            yield text
            return

        for raw_line in resp.iter_lines():
            if stop_event.is_set():
                resp.close()
                return
            if not raw_line:
                continue
            line = raw_line.decode("utf-8") if isinstance(raw_line, bytes) else raw_line
            try:
                data = json.loads(line)
            except json.JSONDecodeError:
                continue
            token = data.get("message", {}).get("content", "")
            if token:
                yield token

    except req_lib.exceptions.Timeout:
        raise
    except Exception as e:
        log.warning(f"_translate_unit_streaming failed: {e}")
        yield text


# ── DOCX translation helpers ─────────────────────────────────────────────────

def _docx_significant_styles(run):
    """Return frozenset of significant formatting tags on a run's rPr."""
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return frozenset()
    SIGNIFICANT = frozenset([
        qn('w:u'), qn('w:strike'), qn('w:dstrike'),
        qn('w:shd'), qn('w:highlight'), qn('w:bdr'),
        qn('w:effectLst'), qn('w:em'),
    ])
    rPr = run.element.rPr
    if rPr is None:
        return frozenset()
    return frozenset(child.tag for child in rPr if child.tag in SIGNIFICANT)


def _docx_is_skippable(run):
    """True for runs that should not be translated (image, tab, instrText, empty)."""
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return False
    xml = getattr(run.element, 'xml', '')
    if '<w:drawing' in xml or '<w:pict' in xml:
        return True
    if run.text == '':
        return True
    if run.text.strip() == '' and ('<w:tab' in xml or '<w:ptab' in xml):
        return True
    if run.element.find(qn('w:instrText')) is not None:
        return True
    return False


def _docx_collect_segments(para):
    """
    Return list of (runs_list, full_text) for each formatting-uniform segment
    in the paragraph. Runs with the same significant styles are grouped together.
    instrText / image / tab runs act as boundaries and are skipped.
    """
    try:
        from docx.oxml.ns import qn
        from docx.text.run import Run
    except ImportError:
        return []

    segments = []
    current_runs = []
    current_styles = None

    for child in para._p:
        tag = child.tag
        # Skip non-run elements (pPr, bookmarks, etc.)
        if not tag.endswith('}r'):
            if current_runs:
                text = ''.join(r.text for r in current_runs)
                if text.strip():
                    segments.append((list(current_runs), text))
                current_runs = []
                current_styles = None
            continue

        run = Run(child, para)

        if _docx_is_skippable(run):
            if current_runs:
                text = ''.join(r.text for r in current_runs)
                if text.strip():
                    segments.append((list(current_runs), text))
                current_runs = []
                current_styles = None
            continue

        styles = _docx_significant_styles(run)
        if current_styles is None:
            current_styles = styles
        elif styles != current_styles:
            text = ''.join(r.text for r in current_runs)
            if text.strip():
                segments.append((list(current_runs), text))
            current_runs = []
            current_styles = styles

        current_runs.append(run)

    if current_runs:
        text = ''.join(r.text for r in current_runs)
        if text.strip():
            segments.append((list(current_runs), text))

    return segments


def _docx_apply_to_segment(runs, translated_text):
    """Write translated_text into the first live run; clear the rest."""
    first = None
    for r in runs:
        if r.element.getparent() is not None:
            first = r
            break
    if first is None:
        return
    first.text = translated_text
    for r in runs[1:]:
        if r.element.getparent() is not None:
            r.text = ''


def _docx_collect_paragraphs(doc):
    """
    Yield all translatable Paragraph objects from body, tables,
    headers/footers, and footnotes/endnotes.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
    from docx.oxml.ns import qn

    def _from_element(parent_elem, container):
        for child in parent_elem:
            tag = child.tag
            if tag.endswith('}p'):
                yield Paragraph(child, container)
            elif tag.endswith('}tbl'):
                tbl = Table(child, container)
                for row in tbl.rows:
                    for cell in row.cells:
                        yield from _from_element(cell._element, cell)
            elif tag.endswith('}sdt'):
                sdt_content = child.find(qn('w:sdtContent'))
                if sdt_content is not None:
                    yield from _from_element(sdt_content, container)

    # Body
    yield from _from_element(doc.element.body, doc)

    # Headers / footers
    for section in doc.sections:
        for hf in (section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer):
            if hf is not None:
                yield from _from_element(hf._element, hf)

    # Footnotes / endnotes
    for attr in ('footnotes_part', 'endnotes_part'):
        part = getattr(doc.part, attr, None)
        if part is not None:
            yield from _from_element(part.element, part)


def translate_docx_bytes(content, base_name, lang_from, lang_to, stop_event):
    """Generator. Yields ('log'|'progress'|'error'|'stopped'|'done', ...)."""
    try:
        from docx import Document
    except ImportError:
        yield ("error", "Бібліотека python-docx не встановлена")
        return

    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        yield ("error", f"Не вдалось відкрити DOCX: {e}")
        return

    # Collect all paragraphs with their formatting segments
    para_jobs = []  # list of (para, segments) where segments = [(runs, text), ...]
    total_chars = 0
    for para in _docx_collect_paragraphs(doc):
        if not para.text.strip():
            continue
        segments = _docx_collect_segments(para)
        if not segments:
            continue
        para_jobs.append((para, segments))
        total_chars += sum(len(t) for _, t in segments)

    total = len(para_jobs)
    yield ("log", f"DOCX: {total} абзаців, {total_chars} символів")

    if total == 0:
        yield ("error", "У файлі не знайдено тексту для перекладу")
        return

    max_chars = CFG.get("max_chars", 30000)
    if total_chars > max_chars:
        yield ("error",
               f"Файл занадто великий: {total_chars} символів "
               f"(максимум {max_chars})")
        return

    yield ("meta", {"chars": total_chars, "pages": None})

    for i, (para, segments) in enumerate(para_jobs, 1):
        if stop_event.is_set():
            yield ("stopped",)
            return

        # Translate the full paragraph text for context, then map back to segments.
        # If the paragraph has only one segment, we translate it directly.
        # If multiple segments (different inline formats), we translate the whole
        # paragraph text and distribute proportionally by segment length.
        full_text = ' '.join(t for _, t in segments)

        try:
            translated_full = _translate_unit(full_text, lang_from, lang_to, stop_event)
        except Exception as e:
            yield ("error", f"Помилка перекладу абзацу {i}: {e}")
            return

        if translated_full is None:
            yield ("stopped",)
            return

        if len(segments) == 1:
            _docx_apply_to_segment(segments[0][0], translated_full)
        else:
            # Distribute translated text across segments proportionally.
            # Split at whitespace boundaries to keep words intact.
            words = translated_full.split()
            total_orig_len = sum(len(t) for _, t in segments)
            word_idx = 0
            for seg_i, (runs, orig_text) in enumerate(segments):
                if seg_i == len(segments) - 1:
                    # Last segment gets all remaining words
                    seg_words = words[word_idx:]
                else:
                    ratio = len(orig_text) / total_orig_len
                    count = max(1, round(len(words) * ratio))
                    seg_words = words[word_idx:word_idx + count]
                    word_idx += count
                seg_text = ' '.join(seg_words) if seg_words else orig_text
                _docx_apply_to_segment(runs, seg_text)

        yield ("progress", f"Абзац {i}/{total}", round(i / total * 100))

    out = io.BytesIO()
    doc.save(out)
    yield ("done",
           f"{base_name}_translated.docx",
           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
           out.getvalue())


def translate_pdf_bytes(content, base_name, lang_from, lang_to, stop_event):
    try:
        import fitz
    except ImportError:
        yield ("error", "Бібліотека PyMuPDF не встановлена")
        return

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        yield ("error", f"Не вдалось відкрити PDF: {e}")
        return

    total_pages = len(doc)
    yield ("log", f"PDF: {total_pages} сторінок")

    max_pages = CFG.get("max_pdf_pages", 10)
    if total_pages > max_pages:
        doc.close()
        yield ("error",
               f"PDF занадто великий: {total_pages} сторінок "
               f"(максимум {max_pages})")
        return

    yield ("meta", {"chars": None, "pages": total_pages})

    for page_num in range(total_pages):
        if stop_event.is_set():
            doc.close()
            yield ("stopped",)
            return

        page = doc[page_num]
        raw = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        blocks = [b for b in raw["blocks"] if b["type"] == 0]

        items = []
        for block in blocks:
            lines_text = []
            fontsize = 11
            color = 0
            for line in block["lines"]:
                for span in line["spans"]:
                    if span["text"].strip():
                        lines_text.append(span["text"])
                        fontsize = span["size"]
                        color = span["color"]
            full_text = " ".join(lines_text).strip()
            if full_text:
                items.append((fitz.Rect(block["bbox"]), full_text, fontsize, color))

        pct = round((page_num + 1) / total_pages * 100)
        if not items:
            yield ("progress", f"Сторінка {page_num + 1}/{total_pages}", pct)
            continue

        for rect, _, _, _ in items:
            page.add_redact_annot(rect, fill=(1, 1, 1))
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        for rect, text, fontsize, color_int in items:
            if stop_event.is_set():
                doc.close()
                yield ("stopped",)
                return

            try:
                translated = _translate_unit(text, lang_from, lang_to, stop_event)
            except Exception as e:
                doc.close()
                yield ("error", f"Помилка перекладу: {e}")
                return

            if translated is None:
                doc.close()
                yield ("stopped",)
                return

            r = ((color_int >> 16) & 0xFF) / 255
            g = ((color_int >> 8) & 0xFF) / 255
            b = (color_int & 0xFF) / 255

            page.insert_textbox(
                rect, translated,
                fontsize=max(fontsize - 0.5, 6),
                color=(r, g, b),
                align=0,
                overflow="ignore",
            )

        yield ("progress", f"Сторінка {page_num + 1}/{total_pages}", pct)

    out = io.BytesIO()
    doc.save(out, garbage=4, deflate=True)
    doc.close()
    yield ("done",
           f"{base_name}_translated.pdf",
           "application/pdf",
           out.getvalue())


def translate_txt_bytes(content, base_name, lang_from, lang_to, stop_event):
    """Per-paragraph TXT translation."""
    try:
        text = content.decode("utf-8", errors="replace")
    except Exception as e:
        yield ("error", f"Не вдалось прочитати файл: {e}")
        return

    max_chars = CFG.get("max_chars", 30000)
    if len(text) > max_chars:
        yield ("error",
               f"Файл занадто великий: {len(text)} символів "
               f"(максимум {max_chars})")
        return

    if not text.strip():
        yield ("error", "Файл порожній")
        return

    paragraphs = text.replace("\r\n", "\n").split("\n\n")
    total = len(paragraphs)
    yield ("log", f"TXT: {total} абзаців, {len(text)} символів")
    yield ("meta", {"chars": len(text), "pages": None})

    translated_parts = []
    for i, para in enumerate(paragraphs, 1):
        if stop_event.is_set():
            yield ("stopped",)
            return

        if not para.strip():
            translated_parts.append(para)
            continue

        try:
            result = _translate_unit(para, lang_from, lang_to, stop_event)
        except Exception as e:
            yield ("error", str(e))
            return

        if result is None:
            yield ("stopped",)
            return

        translated_parts.append(result)
        yield ("progress", f"Абзац {i}/{total}", round(i / total * 100))

    final = "\n\n".join(translated_parts)
    yield ("done",
           f"{base_name}_translated.txt",
           "text/plain; charset=utf-8",
           final.encode("utf-8"))


USER_HTML = r"""<!DOCTYPE html>
<html lang="uk">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Перекладач</title>
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  :root {
    --bg: #f8fafc;
    --card: #ffffff;
    --border: #e2e8f0;
    --text: #1e293b;
    --muted: #64748b;
    --primary: #2563eb;
    --primary-hover: #1d4ed8;
    --primary-light: #dbeafe;
    --danger: #dc2626;
    --danger-hover: #b91c1c;
    --success: #16a34a;
    --success-hover: #15803d;
    --shadow: 0 1px 3px rgba(0,0,0,0.04), 0 1px 2px rgba(0,0,0,0.06);
  }
  body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
    background: var(--bg);
    color: var(--text);
    padding: 32px 20px;
    min-height: 100vh;
    line-height: 1.5;
  }
  .container { width: 80%; max-width: 100%; margin: 0 auto; }
  header { text-align: center; margin-bottom: 32px; }
  header h1 { font-size: 1.75rem; font-weight: 600; color: var(--text); }
  header p { color: var(--muted); font-size: 0.95rem; margin-top: 4px; }

  .card {
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 24px;
    box-shadow: var(--shadow);
    margin-bottom: 16px;
  }

  .lang-bar {
    display: grid;
    grid-template-columns: 1fr auto 1fr;
    gap: 12px;
    align-items: end;
  }
  .lang-bar label { display: block; font-size: 0.8rem; color: var(--muted); margin-bottom: 6px; font-weight: 500; }
  .lang-bar select {
    width: 100%; padding: 10px 12px; border: 1px solid var(--border); border-radius: 8px;
    font-size: 0.95rem; background: var(--card); cursor: pointer;
  }
  .lang-bar select:focus { outline: none; border-color: var(--primary); box-shadow: 0 0 0 3px var(--primary-light); }
  .swap-btn {
    width: 40px; height: 40px; border-radius: 8px; border: 1px solid var(--border);
    background: var(--card); cursor: pointer; font-size: 1.1rem; color: var(--muted);
    display: flex; align-items: center; justify-content: center;
    transition: all 0.15s;
  }
  .swap-btn:hover { border-color: var(--primary); color: var(--primary); }

  .tabs { display: flex; gap: 4px; margin-bottom: 16px; background: var(--border); padding: 4px; border-radius: 10px; }
  .tab {
    flex: 1; padding: 10px; border: none; background: transparent; cursor: pointer;
    border-radius: 7px; font-size: 0.95rem; font-weight: 500; color: var(--muted);
    transition: all 0.15s;
  }
  .tab.active { background: var(--card); color: var(--text); box-shadow: var(--shadow); }

  .panel { display: none; }
  .panel.active { display: block; }

  textarea {
    width: 100%; padding: 14px; border: 1px solid var(--border); border-radius: 8px;
    font-size: 0.95rem; font-family: inherit; resize: vertical; min-height: 140px;
    line-height: 1.6; background: var(--card); color: var(--text);
  }
  textarea:focus { outline: none; border-color: var(--primary); box-shadow: 0 0 0 3px var(--primary-light); }
  textarea[readonly] { background: #f1f5f9; }

  .rich-editor {
    width: 100%; padding: 14px; border: 1px solid var(--border); border-radius: 8px;
    font-size: 0.95rem; font-family: inherit; min-height: 140px; max-height: 400px;
    line-height: 1.6; background: var(--card); color: var(--text);
    overflow-y: auto; overflow-x: hidden; box-sizing: border-box; outline: none;
    word-break: break-word; overflow-wrap: break-word;
  }
  .rich-editor:focus { border-color: var(--primary); box-shadow: 0 0 0 3px var(--primary-light); }
  .rich-editor[contenteditable=true]:empty:before {
    content: attr(data-placeholder); color: #aaa; pointer-events: none;
  }
  .rich-editor b, .rich-editor strong { font-weight: bold; }
  .rich-editor i, .rich-editor em { font-style: italic; }
  .rich-editor u { text-decoration: underline; }
  .rich-editor h1, .rich-editor h2, .rich-editor h3 { margin: 0.3em 0; font-weight: bold; }
  .rich-editor h1 { font-size: 1.4em; }
  .rich-editor h2 { font-size: 1.2em; }
  .rich-editor h3 { font-size: 1.05em; }
  .rich-editor ul, .rich-editor ol { margin: 4px 0 4px 20px; padding: 0; }
  .rich-editor table { border-collapse: collapse; width: 100%; margin: 4px 0; }
  .rich-editor td, .rich-editor th { border: 1px solid var(--border); padding: 4px 8px; }

  .result-content {
    width: 100%; padding: 14px; border: 1px solid var(--border); border-radius: 8px;
    font-size: 0.95rem; font-family: inherit; min-height: 80px;
    line-height: 1.6; background: #f1f5f9; color: var(--text);
    box-sizing: border-box; overflow-y: auto; overflow-x: hidden;
    word-break: break-word; overflow-wrap: break-word;
  }
  .result-content b, .result-content strong { font-weight: bold; }
  .result-content i, .result-content em { font-style: italic; }
  .result-content u { text-decoration: underline; }
  .result-content h1, .result-content h2, .result-content h3 { margin: 0.3em 0; font-weight: bold; }
  .result-content h1 { font-size: 1.4em; }
  .result-content h2 { font-size: 1.2em; }
  .result-content h3 { font-size: 1.05em; }
  .result-content ul, .result-content ol { margin: 4px 0 4px 20px; padding: 0; }
  .result-content table { border-collapse: collapse; width: 100%; margin: 4px 0; }
  .result-content td, .result-content th { border: 1px solid var(--border); padding: 4px 8px; }

  .drop-zone {
    border: 2px dashed var(--border); border-radius: 10px; padding: 40px 20px;
    text-align: center; cursor: pointer; transition: all 0.15s; background: var(--card);
  }
  .drop-zone:hover, .drop-zone.drag-over { border-color: var(--primary); background: var(--primary-light); }
  .drop-zone-icon { font-size: 2rem; margin-bottom: 8px; }
  .drop-zone-text { color: var(--muted); font-size: 0.95rem; }
  .drop-zone-hint { color: var(--muted); font-size: 0.8rem; margin-top: 4px; }
  .drop-zone input { display: none; }
  .file-name { margin-top: 12px; padding: 10px 12px; background: var(--primary-light); border-radius: 8px; font-size: 0.9rem; color: var(--primary); }

  .actions { display: flex; gap: 10px; margin-top: 14px; align-items: center; flex-wrap: wrap; }
  button.primary, .download-link {
    background: var(--primary); color: white; border: none; border-radius: 8px;
    padding: 11px 22px; font-size: 0.95rem; font-weight: 500; cursor: pointer;
    transition: background 0.15s; text-decoration: none; display: inline-flex; align-items: center; gap: 6px;
  }
  button.primary:hover { background: var(--primary-hover); }
  button.primary:disabled { background: #93c5fd; cursor: not-allowed; }
  button.stop {
    background: var(--danger); color: white; border: none; border-radius: 8px;
    padding: 11px 22px; font-size: 0.95rem; font-weight: 500; cursor: pointer;
    transition: background 0.15s; display: none;
  }
  button.stop:hover { background: var(--danger-hover); }
  button.stop.visible { display: inline-block; }
  .download-link { background: var(--success); display: none; }
  .download-link:hover { background: var(--success-hover); }
  .download-link.visible { display: inline-flex; }

  .status { font-size: 0.85rem; color: var(--muted); display: flex; align-items: center; gap: 8px; }
  .status.error { color: var(--danger); }
  .status.success { color: var(--success); }
  .spinner {
    width: 14px; height: 14px; border: 2px solid var(--border); border-top-color: var(--primary);
    border-radius: 50%; animation: spin 0.8s linear infinite; display: none;
  }
  .spinner.visible { display: inline-block; }
  @keyframes spin { to { transform: rotate(360deg); } }

  .result-label { font-size: 0.85rem; color: var(--muted); margin-bottom: 8px; font-weight: 500; }
  .footer { text-align: center; margin-top: 24px; font-size: 0.8rem; color: var(--muted); }
  .footer a { color: var(--muted); text-decoration: none; }
  .footer a:hover { color: var(--primary); }

  /* Side-by-side text translation layout */
  .split-layout {
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 16px;
    align-items: start;
  }
  @media (max-width: 640px) {
    .split-layout { grid-template-columns: 1fr; }
  }
  .split-panel { display: flex; flex-direction: column; min-width: 0; overflow: hidden; }
  .panel-label {
    font-size: 0.8rem; color: var(--muted); font-weight: 500;
    margin-bottom: 8px; display: flex; align-items: center; gap: 8px; min-height: 22px;
  }
  .rich-editor { min-height: 320px; max-height: 70vh; }
  .result-content { min-height: 320px; max-height: 70vh; }
  .result-content.streaming { color: var(--muted); font-style: italic; white-space: pre-wrap; }
  .result-placeholder { color: #aaa; font-size: 0.9rem; padding: 14px; }

  .model-bar { display: flex; align-items: center; gap: 16px; margin-top: 12px; flex-wrap: wrap; }
  .conn-status { display: flex; align-items: center; gap: 6px; font-size: 0.82rem; color: var(--muted); white-space: nowrap; }
  .conn-dot { width: 9px; height: 9px; border-radius: 50%; flex-shrink: 0; background: #94a3b8; transition: background 0.3s; }
  .conn-dot.ok { background: var(--success); }
  .conn-dot.err { background: var(--danger); }
  .model-select-wrap { flex: 1; min-width: 200px; }
  .model-select-wrap label { display: block; font-size: 0.78rem; color: var(--muted); margin-bottom: 4px; }
  .model-select-wrap select { width: 100%; }
</style>
</head>
<body>
<div class="container">
  <header>
    <h1>Перекладач</h1>
    <p>Локальний переклад документів і тексту</p>
  </header>

  <div class="card">
    <div class="lang-bar">
      <div>
        <label>З мови</label>
        <select id="lang_from"></select>
      </div>
      <button class="swap-btn" onclick="swapLangs()" title="Поміняти мови місцями">⇄</button>
      <div>
        <label>На мову</label>
        <select id="lang_to"></select>
      </div>
    </div>
    <div class="model-bar">
      <div class="conn-status">
        <span class="conn-dot" id="conn-dot"></span>
        <span id="conn-text">Перевірка...</span>
      </div>
      <div class="model-select-wrap">
        <label>Модель</label>
        <select id="model_select" onchange="onModelChange()"></select>
      </div>
    </div>
  </div>

  <div class="card">
    <div class="tabs">
      <button class="tab active" onclick="showTab('text')">Текст</button>
      <button class="tab" onclick="showTab('file')">Файл</button>
    </div>

    <div id="panel-text" class="panel active">
      <div class="split-layout">
        <div class="split-panel">
          <div class="panel-label">Оригінал</div>
          <div id="input" class="rich-editor" contenteditable="true" spellcheck="false"
               data-gramm="false" data-placeholder="Введіть текст для перекладу..."></div>
          <div class="actions">
            <button id="btn-translate" class="primary" onclick="doTranslate()">Перекласти</button>
            <button id="btn-stop" class="stop" onclick="doStop()">Зупинити</button>
            <span class="spinner" id="text-spinner"></span>
            <span class="status" id="text-status"></span>
          </div>
        </div>
        <div class="split-panel">
          <div class="panel-label" id="result-label">Переклад</div>
          <div id="result" class="result-content">
            <div class="result-placeholder">Переклад з'явиться тут...</div>
          </div>
        </div>
      </div>
    </div>

    <div id="panel-file" class="panel">
      <div class="drop-zone" id="drop-zone" onclick="document.getElementById('file-input').click()">
        <input type="file" id="file-input" accept=".pdf,.docx,.txt" onchange="fileSelected(this.files[0])">
        <div class="drop-zone-icon">📄</div>
        <div class="drop-zone-text">Натисніть або перетягніть файл</div>
        <div class="drop-zone-hint">Підтримуються DOCX, PDF, TXT</div>
      </div>
      <div class="file-name" id="file-name" style="display:none"></div>
      <div class="actions">
        <button id="btn-file" class="primary" onclick="doTranslateFile()">Перекласти файл</button>
        <button id="btn-file-stop" class="stop" onclick="doFileStop()">Зупинити</button>
        <a id="download-link" class="download-link">↓ Скачати результат</a>
        <span class="spinner" id="file-spinner"></span>
        <span class="status" id="file-status"></span>
      </div>
    </div>
  </div>

  <div class="footer">
    <a href="/admin">Адміністрування</a>
    <span style="margin: 0 12px; color: var(--border);">|</span>
    <span>v1.2</span>
  </div>
</div>

<script>
// ── Tabs ──────────────────────────────────────────────────────────────
function showTab(name) {
  document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
  document.querySelectorAll('.panel').forEach(p => p.classList.remove('active'));
  event.target.classList.add('active');
  document.getElementById('panel-' + name).classList.add('active');
}

// ── Languages ─────────────────────────────────────────────────────────
async function initLangs() {
  const r = await fetch('/languages');
  const langs = await r.json();
  const from = document.getElementById('lang_from');
  const to = document.getElementById('lang_to');
  from.appendChild(new Option('Автовизначення', 'auto'));
  langs.forEach(({label, value}) => {
    from.appendChild(new Option(label, value));
    to.appendChild(new Option(label, value));
  });
  to.value = 'Ukrainian';
}

function swapLangs() {
  const from = document.getElementById('lang_from');
  const to = document.getElementById('lang_to');
  if (from.value === 'auto') return;
  const tmp = from.value;
  from.value = to.value;
  to.value = tmp;
}

// ── Status helpers ────────────────────────────────────────────────────
function setStatus(id, text, type='') {
  const el = document.getElementById(id);
  el.textContent = text;
  el.className = 'status' + (type ? ' ' + type : '');
}
function setSpinner(id, on) {
  document.getElementById(id).classList.toggle('visible', on);
}
function setWorking(prefix, on) {
  document.getElementById('btn-' + (prefix === 'text' ? 'translate' : 'file')).disabled = on;
  document.getElementById('btn-' + (prefix === 'text' ? 'stop' : 'file-stop')).classList.toggle('visible', on);
  setSpinner(prefix + '-spinner', on);
}

// ── Text translation ──────────────────────────────────────────────────
let _textController = null;
let _textRequestId = null;
let _textTokens = '';

async function doTranslate() {
  const inputEl = document.getElementById('input');
  const html = inputEl.innerHTML.trim();
  const text = inputEl.textContent.trim();
  if (!text) { setStatus('text-status', 'Введіть текст', 'error'); return; }

  setWorking('text', true);
  setStatus('text-status', 'Перекладаю...');
  document.getElementById('result').innerHTML = '';
  _textTokens = '';
  _textController = new AbortController();

  try {
    const resp = await fetch('/translate-html', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      signal: _textController.signal,
      body: JSON.stringify({
        html,
        lang_from: document.getElementById('lang_from').value,
        lang_to: document.getElementById('lang_to').value,
      }),
    });
    const reader = resp.body.getReader();
    const decoder = new TextDecoder();
    let buffer = '';
    while (true) {
      const {done, value} = await reader.read();
      if (done) break;
      buffer += decoder.decode(value, {stream: true});
      const lines = buffer.split('\n');
      buffer = lines.pop();
      for (const line of lines) {
        if (!line.startsWith('data: ')) continue;
        let evt; try { evt = JSON.parse(line.slice(6)); } catch { continue; }
        if (evt.type === 'id') _textRequestId = evt.text;
        else if (evt.type === 'queue' && evt.ahead > 0) setStatus('text-status', `У черзі: попереду ${evt.ahead}`);
        else if (evt.type === 'queue' && evt.ahead === 0) setStatus('text-status', 'Перекладаю...');
        else if (evt.type === 'token') {
          // Stream tokens into result as plain text (fast visual feedback)
          _textTokens += evt.text;
          const resultEl = document.getElementById('result');
          resultEl.classList.add('streaming');
          resultEl.textContent = _textTokens;
        }
        else if (evt.type === 'result') {
          showResult(evt.text, evt.format);
          setStatus('text-status', 'Готово', 'success');
        }
        else if (evt.type === 'error') {
          setStatus('text-status', 'Помилка: ' + (evt.text || 'невідома'), 'error');
        }
      }
    }
  } catch (e) {
    if (e.name !== 'AbortError') setStatus('text-status', 'Помилка з\'єднання', 'error');
  }
  setWorking('text', false);
  _textController = null;
  _textRequestId = null;
}

function showResult(text, format) {
  const resultEl = document.getElementById('result');
  resultEl.classList.remove('streaming');
  if (format === 'markdown' && typeof marked !== 'undefined') {
    resultEl.innerHTML = marked.parse(text);
  } else {
    resultEl.innerHTML = text;
  }
}

document.getElementById('input').addEventListener('keydown', function(e) {
  if (e.ctrlKey && e.key === 'Enter') doTranslate();
});

async function doStop() {
  if (_textController) { _textController.abort(); _textController = null; }
  if (_textRequestId) {
    await fetch('/stop/' + _textRequestId, {method: 'POST'}).catch(() => {});
    _textRequestId = null;
  }
  setStatus('text-status', 'Зупинено');
  setWorking('text', false);
}

// ── File translation ──────────────────────────────────────────────────
let _fileController = null;
let _fileRequestId = null;
let _selectedFile = null;

function fileSelected(f) {
  _selectedFile = f;
  const nameEl = document.getElementById('file-name');
  if (f) {
    nameEl.textContent = f.name + ' (' + (f.size / 1024).toFixed(1) + ' KB)';
    nameEl.style.display = 'block';
  } else {
    nameEl.style.display = 'none';
  }
  document.getElementById('download-link').classList.remove('visible');
  setStatus('file-status', '');
}

const dropZone = document.getElementById('drop-zone');
dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('drag-over'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('drag-over'));
dropZone.addEventListener('drop', e => {
  e.preventDefault();
  dropZone.classList.remove('drag-over');
  if (e.dataTransfer.files[0]) fileSelected(e.dataTransfer.files[0]);
});

async function doTranslateFile() {
  if (!_selectedFile) { setStatus('file-status', 'Оберіть файл', 'error'); return; }

  setWorking('file', true);
  setStatus('file-status', 'Завантаження...');
  document.getElementById('download-link').classList.remove('visible');
  _fileController = new AbortController();

  const fd = new FormData();
  fd.append('file', _selectedFile, _selectedFile.name);
  fd.append('lang_from', document.getElementById('lang_from').value);
  fd.append('lang_to', document.getElementById('lang_to').value);

  try {
    const resp = await fetch('/translate-file', {
      method: 'POST',
      signal: _fileController.signal,
      body: fd,
    });
    const reader = resp.body.getReader();
    const decoder = new TextDecoder();
    let buffer = '';
    while (true) {
      const {done, value} = await reader.read();
      if (done) break;
      buffer += decoder.decode(value, {stream: true});
      const lines = buffer.split('\n');
      buffer = lines.pop();
      for (const line of lines) {
        if (!line.startsWith('data: ')) continue;
        let evt; try { evt = JSON.parse(line.slice(6)); } catch { continue; }
        if (evt.type === 'id') _fileRequestId = evt.text;
        else if (evt.type === 'progress') setStatus('file-status', evt.text);
        else if (evt.type === 'download') {
          const link = document.getElementById('download-link');
          link.href = evt.url;
          link.download = evt.filename;
          link.classList.add('visible');
          setStatus('file-status', 'Готово', 'success');
        }
        else if (evt.type === 'error') {
          setStatus('file-status', 'Помилка: ' + (evt.text || 'невідома'), 'error');
        }
      }
    }
  } catch (e) {
    if (e.name !== 'AbortError') setStatus('file-status', 'Помилка з\'єднання', 'error');
  }
  setWorking('file', false);
  _fileController = null;
  _fileRequestId = null;
}

async function doFileStop() {
  if (_fileController) { _fileController.abort(); _fileController = null; }
  if (_fileRequestId) {
    await fetch('/stop/' + _fileRequestId, {method: 'POST'}).catch(() => {});
    _fileRequestId = null;
  }
  setStatus('file-status', 'Зупинено');
  setWorking('file', false);
}

// ── Init ──────────────────────────────────────────────────────────────
initLangs();
document.getElementById('input').addEventListener('keydown', e => {
  if (e.ctrlKey && e.key === 'Enter') doTranslate();
});
</script>
</body>
</html>"""


ADMIN_HTML = r"""<!DOCTYPE html>
<html lang="uk">
<head>
<meta charset="UTF-8">
<title>Перекладач — Адмін</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: sans-serif; background: #f5f5f5; padding: 24px; max-width: 900px; margin: 0 auto; }
  h1 { margin-bottom: 20px; font-size: 1.4rem; color: #333; }
  .card { background: white; border-radius: 8px; padding: 20px; box-shadow: 0 1px 4px rgba(0,0,0,.1); margin-bottom: 16px; }
  label { display: block; font-size: .85rem; color: #555; margin-bottom: 4px; }
  select, textarea, input[type=text], input[type=number] { width: 100%; border: 1px solid #ddd; border-radius: 6px; padding: 8px 10px; font-size: .95rem; font-family: inherit; }
  textarea { resize: vertical; }
  .row { display: flex; gap: 12px; margin-bottom: 12px; }
  .row > div { flex: 1; }
  .btn-row { display: flex; gap: 8px; align-items: center; flex-wrap: wrap; }
  button {
    background: #2563eb; color: white; border: none; border-radius: 6px;
    padding: 10px 24px; font-size: 1rem; cursor: pointer; transition: background .2s;
  }
  button:hover { background: #1d4ed8; }
  button:disabled { background: #93c5fd; cursor: not-allowed; }
  .btn-stop { background: #dc2626; display: none; }
  .btn-stop:hover { background: #b91c1c; }
  .btn-save { background: #16a34a; }
  .btn-save:hover { background: #15803d; }
  .btn-download { background: #16a34a; text-decoration: none; display: none; padding: 10px 20px; border-radius: 6px; font-size: 1rem; color: white; }
  .btn-download:hover { background: #15803d; }
  .log-box {
    background: #1e1e1e; color: #d4d4d4; font-family: monospace; font-size: .8rem;
    padding: 12px; border-radius: 6px; height: 160px; overflow-y: auto;
    white-space: pre-wrap; word-break: break-all; margin-top: 12px;
  }
  .status { font-size: .8rem; color: #888; margin-top: 6px; }
  .status.ok { color: #16a34a; }
  .status.err { color: #dc2626; }
  details { margin-top: 12px; }
  summary { font-size: .85rem; color: #555; cursor: pointer; user-select: none; }
  summary:hover { color: #333; }
  .hint { font-size: .75rem; color: #aaa; margin-top: 4px; }
  .file-drop {
    border: 2px dashed #ddd; border-radius: 6px; padding: 24px; text-align: center;
    color: #aaa; font-size: .9rem; cursor: pointer; transition: border-color .2s, color .2s;
  }
  .file-drop.drag-over { border-color: #2563eb; color: #2563eb; }
  .file-drop input { display: none; }
  .file-name { font-size: .85rem; color: #555; margin-top: 6px; }
  .progress-bar-wrap { height: 6px; background: #e5e7eb; border-radius: 3px; margin-top: 10px; display: none; }
  .progress-bar { height: 100%; background: #2563eb; border-radius: 3px; width: 0; transition: width .3s; }
  .settings-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-top: 12px; }
  .settings-grid .full { grid-column: 1 / -1; }
  .section-title { font-size: 1rem; font-weight: 600; color: #333; margin-bottom: 12px; }
</style>
</head>
<body>
<div style="background:#fef9c3; border:1px solid #fde047; border-radius:8px; padding:14px 20px; margin-bottom:20px; display:flex; align-items:center; gap:12px;">
  <span style="font-size:1.4rem;">🧪</span>
  <div>
    <strong style="color:#854d0e;">Тестовий режим</strong>
    <span style="color:#713f12; font-size:.9rem;"> — сервіс працює в режимі тестування. Обробка одного запиту може займати до <strong>5 хвилин</strong>. Дякуємо за терпіння!</span>
  </div>
</div>
<h1>LocalAI Перекладач</h1>

<!-- Language selectors + Advanced (shared) -->
<div class="card">
  <div class="row">
    <div>
      <label>Перекласти з</label>
      <select id="lang_from"></select>
    </div>
    <div>
      <label>Перекласти на</label>
      <select id="lang_to"></select>
    </div>
  </div>

</div>

<!-- Text translation -->
<div class="card">
  <label style="margin-bottom:8px; font-size:1rem; color:#333; font-weight:600;">Переклад тексту</label>
  <div style="margin-bottom:12px; margin-top:8px;">
    <textarea id="input" rows="6" placeholder="Введіть текст тут..."></textarea>
  </div>
  <div id="queue-banner" style="display:none; background:#fff7ed; border:1px solid #fb923c; border-radius:8px; padding:12px 16px; margin-bottom:12px; gap:10px;">
    <span style="font-size:1.3rem;">⏳</span>
    <div>
      <strong style="color:#9a3412;">Запит у черзі</strong>
      <span id="queue-banner-text" style="color:#7c2d12; font-size:.9rem;"></span>
    </div>
  </div>
  <div class="btn-row">
    <button id="btn" onclick="doTranslate()">Перекласти</button>
    <button id="stopBtn" class="btn-stop" onclick="doStop()">Зупинити</button>
    <span class="status" id="status"></span>
  </div>
</div>

<div class="card">
  <label style="margin-bottom:8px">Результат перекладу</label>
  <textarea id="result" rows="6" readonly placeholder="Результат з'явиться тут..."></textarea>
</div>

<div class="card" id="thinking-card" style="display:none;">
  <details id="thinking-details">
    <summary style="font-size:.85rem; color:#7c3aed; cursor:pointer; user-select:none;">&#129504; Думки моделі</summary>
    <div id="thinking" style="background:#faf5ff; border:1px solid #e9d5ff; border-radius:6px; padding:10px; margin-top:8px; font-family:monospace; font-size:.78rem; color:#4c1d95; white-space:pre-wrap; word-break:break-word; max-height:300px; overflow-y:auto;"></div>
  </details>
</div>

<div class="card">
  <label style="margin-bottom:8px">Лог</label>
  <div id="log" class="log-box"></div>
</div>

<!-- File translation -->
<div class="card">
  <label style="margin-bottom:8px; font-size:1rem; color:#333; font-weight:600;">Переклад файлу</label>
  <div style="margin-top:8px; margin-bottom:12px;">
    <div class="file-drop" id="file-drop" onclick="document.getElementById('file-input').click()"
         ondragover="fileDragOver(event)" ondragleave="fileDragLeave(event)" ondrop="fileDrop(event)">
      <input type="file" id="file-input" accept=".pdf,.txt,.docx" onchange="fileSelected(this)">
      Натисніть або перетягніть файл сюди<br>
      <span style="font-size:.75rem">PDF, DOCX, TXT</span>
    </div>
    <div class="file-name" id="file-name"></div>
    <div class="progress-bar-wrap" id="progress-wrap"><div class="progress-bar" id="progress-bar"></div></div>
  </div>
  <div class="btn-row">
    <button id="file-btn" onclick="doTranslateFile()">Перекласти файл</button>
    <button id="file-stopBtn" class="btn-stop" onclick="doFileStop()">Зупинити</button>
    <a id="download-link" class="btn-download">&#8595; Завантажити</a>
    <span class="status" id="file-status"></span>
  </div>
  <div id="file-log" class="log-box" style="display:none;"></div>
</div>

<!-- Stats -->
<div class="card">
  <details id="stats-details">
    <summary class="section-title" style="margin-bottom:0">&#128202; Статистика</summary>
    <div style="margin-top:12px;">
      <div id="stats-summary" style="display:grid; grid-template-columns: repeat(auto-fit, minmax(120px, 1fr)); gap:10px; margin-bottom:14px;"></div>
      <div style="overflow-x:auto;">
        <table id="stats-table" style="width:100%; border-collapse: collapse; font-size: .82rem;">
          <thead>
            <tr style="background:#f1f5f9; text-align:left;">
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Час</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">IP</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Тип</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Файл</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Мови</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Симв.</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Стор.</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Час, с</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Статус</th>
              <th style="padding:6px 8px; border-bottom:1px solid #e2e8f0;">Помилка</th>
            </tr>
          </thead>
          <tbody id="stats-tbody"></tbody>
        </table>
      </div>
      <div class="btn-row" style="margin-top:10px;">
        <button onclick="loadStats()" style="background:#6b7280;">&#8635; Оновити</button>
      </div>
    </div>
  </details>
</div>

<!-- Settings -->
<div class="card">
  <details id="settings-details">
    <summary class="section-title" style="margin-bottom:0">&#9881; Налаштування</summary>
    <div class="settings-grid">
      <div class="full">
        <label>LocalAI URL</label>
        <input type="text" id="cfg_base_url" placeholder="http://192.168.x.x:port/v1">
      </div>
      <div class="full">
        <label>Модель</label>
        <input type="text" id="cfg_model" placeholder="назва моделі як у LocalAI">
      </div>
      <div>
        <label>Макс. токенів</label>
        <input type="number" id="cfg_max_tokens" min="128" max="32000" step="128">
      </div>
      <div>
        <label>Таймаут LLM (секунди)</label>
        <input type="number" id="cfg_llm_timeout" min="10" max="600">
      </div>
      <div>
        <label>Розмір чанка (символів)</label>
        <input type="number" id="cfg_chunk_size" min="500" max="20000" step="100">
      </div>
      <div>
        <label>Макс. сторінок PDF</label>
        <input type="number" id="cfg_max_pdf_pages" min="1" max="500" step="1">
      </div>
      <div>
        <label>Макс. символів (текст/DOCX/TXT)</label>
        <input type="number" id="cfg_max_chars" min="1000" max="1000000" step="1000">
      </div>
      <div></div>
    </div>
    <div class="btn-row" style="margin-top:12px">
      <button class="btn-save" onclick="saveSettings()">Зберегти</button>
      <button onclick="resetSettings()" style="background:#6b7280">Скинути до стандартних</button>
      <span class="status" id="cfg-status"></span>
    </div>
  </details>
</div>

<script>
// ── Settings ──────────────────────────────────────────────────────────
async function loadSettings() {
  const r = await fetch('/config');
  const cfg = await r.json();
  document.getElementById('cfg_base_url').value        = cfg.base_url        ?? '';
  document.getElementById('cfg_model').value           = cfg.model           ?? '';
  document.getElementById('cfg_max_tokens').value      = cfg.max_tokens      ?? 2048;
  document.getElementById('cfg_llm_timeout').value     = cfg.llm_timeout     ?? 180;
  document.getElementById('cfg_chunk_size').value      = cfg.chunk_size      ?? 3000;
  document.getElementById('cfg_max_pdf_pages').value   = cfg.max_pdf_pages   ?? 10;
  document.getElementById('cfg_max_chars').value       = cfg.max_chars       ?? 30000;
}

async function saveSettings() {
  const cfg = {
    base_url:        document.getElementById('cfg_base_url').value.trim(),
    model:           document.getElementById('cfg_model').value.trim(),
    max_tokens:      parseInt(document.getElementById('cfg_max_tokens').value),
    llm_timeout:     parseInt(document.getElementById('cfg_llm_timeout').value),
    chunk_size:      parseInt(document.getElementById('cfg_chunk_size').value),
    max_pdf_pages:   parseInt(document.getElementById('cfg_max_pdf_pages').value),
    max_chars:       parseInt(document.getElementById('cfg_max_chars').value),
  };
  const st = document.getElementById('cfg-status');
  try {
    const r = await fetch('/config', { method: 'POST', headers: {'Content-Type':'application/json'}, body: JSON.stringify(cfg) });
    const res = await r.json();
    if (res.ok) { st.textContent = 'Збережено'; st.className = 'status ok'; }
    else        { st.textContent = 'Помилка: ' + (res.error ?? 'невідома'); st.className = 'status err'; }
  } catch(e) { st.textContent = 'Помилка: ' + e; st.className = 'status err'; }
  setTimeout(() => { st.textContent = ''; st.className = 'status'; }, 3000);
}

async function resetSettings() {
  const r = await fetch('/config/defaults');
  const cfg = await r.json();
  document.getElementById('cfg_base_url').value        = cfg.base_url;
  document.getElementById('cfg_model').value           = cfg.model;
  document.getElementById('cfg_max_tokens').value      = cfg.max_tokens;
  document.getElementById('cfg_llm_timeout').value     = cfg.llm_timeout;
  document.getElementById('cfg_chunk_size').value      = cfg.chunk_size;
  document.getElementById('cfg_max_pdf_pages').value   = cfg.max_pdf_pages;
  document.getElementById('cfg_max_chars').value       = cfg.max_chars;
}

// ── Text translation ──────────────────────────────────────────────────
let _controller = null;
let _requestId = null;
let _thinkMode = false;
let _tagBuf = '';

function handleToken(text) {
  const resultEl = document.getElementById('result');
  const thinkEl  = document.getElementById('thinking');
  const thinkCard = document.getElementById('thinking-card');
  _tagBuf += text;
  while (_tagBuf.length > 0) {
    if (_thinkMode) {
      const end = _tagBuf.indexOf('</think>');
      if (end !== -1) {
        thinkEl.textContent += _tagBuf.slice(0, end);
        thinkEl.scrollTop = thinkEl.scrollHeight;
        _tagBuf = _tagBuf.slice(end + 8);
        _thinkMode = false;
      } else if (_tagBuf.length > 8) {
        const safe = _tagBuf.slice(0, _tagBuf.length - 8);
        thinkEl.textContent += safe;
        thinkEl.scrollTop = thinkEl.scrollHeight;
        _tagBuf = _tagBuf.slice(safe.length);
        break;
      } else { break; }
    } else {
      const start = _tagBuf.indexOf('<think>');
      if (start !== -1) {
        resultEl.value += _tagBuf.slice(0, start);
        _tagBuf = _tagBuf.slice(start + 7);
        _thinkMode = true;
        thinkCard.style.display = 'block';
        document.getElementById('thinking-details').open = true;
      } else if (_tagBuf.length > 7) {
        const safe = _tagBuf.slice(0, _tagBuf.length - 7);
        resultEl.value += safe;
        _tagBuf = _tagBuf.slice(safe.length);
        break;
      } else { break; }
    }
  }
}

function setWorking(on) {
  document.getElementById('btn').disabled = on;
  document.getElementById('stopBtn').style.display = on ? 'inline-block' : 'none';
}

async function doStop() {
  if (_controller) { _controller.abort(); _controller = null; }
  if (_requestId) {
    await fetch('/stop/' + _requestId, { method: 'POST' }).catch(() => {});
    _requestId = null;
  }
  document.getElementById('queue-banner').style.display = 'none';
  document.getElementById('status').textContent = 'Зупинено';
  setWorking(false);
}

async function doTranslate() {
  const text = document.getElementById('input').value.trim();
  if (!text) { document.getElementById('status').textContent = 'ПОМИЛКА: текст порожній'; return; }

  const log = document.getElementById('log');
  const result = document.getElementById('result');
  const status = document.getElementById('status');

  setWorking(true);
  document.getElementById('queue-banner').style.display = 'none';
  log.textContent = '';
  result.value = '';
  status.textContent = 'Перекладаю...';
  document.getElementById('thinking').textContent = '';
  document.getElementById('thinking-card').style.display = 'none';
  _thinkMode = false;
  _tagBuf = '';

  const startTime = Date.now();
  _controller = new AbortController();

  try {
    const response = await fetch('/translate', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      signal: _controller.signal,
      body: JSON.stringify({
        text,
        lang_from: document.getElementById('lang_from').value,
        lang_to: document.getElementById('lang_to').value,
      }),
    });

    const reader = response.body.getReader();
    const decoder = new TextDecoder();
    let buffer = '';
    while (true) {
      const { done, value } = await reader.read();
      if (done) break;
      buffer += decoder.decode(value, { stream: true });
      const lines = buffer.split('\n');
      buffer = lines.pop();
      for (const line of lines) {
        if (!line.startsWith('data: ')) continue;
        let evt; try { evt = JSON.parse(line.slice(6)); } catch { continue; }
        if (evt.type === 'id') { _requestId = evt.text; }
        else if (evt.type === 'thinking') {
          const thinkingCard = document.getElementById('thinking-card');
          const thinkingBox = document.getElementById('thinking');
          const details = document.getElementById('thinking-details');
          thinkingCard.style.display = 'block';
          details.open = true;
          thinkingBox.textContent += evt.text;
          thinkingBox.scrollTop = thinkingBox.scrollHeight;
        }
        else if (evt.type === 'queue') {
          const banner = document.getElementById('queue-banner');
          const bannerText = document.getElementById('queue-banner-text');
          if (evt.ahead > 0) {
            bannerText.textContent = ` — попереду ${evt.ahead} запит${evt.ahead > 1 ? 'и' : ''}. Очікуйте...`;
            banner.style.cssText = 'display:flex; align-items:center; gap:10px; background:#fff7ed; border:1px solid #fb923c; border-radius:8px; padding:12px 16px; margin-bottom:12px;';
            status.textContent = '';
          } else {
            banner.style.display = 'none';
            status.textContent = 'Перекладаю...';
          }
        }
        else if (evt.type === 'log') { log.textContent += evt.text + '\n'; log.scrollTop = log.scrollHeight; }
        else if (evt.type === 'token') { handleToken(evt.text); }
        else if (evt.type === 'result') {
          result.value = evt.text;
          status.textContent = `Готово за ${((Date.now()-startTime)/1000).toFixed(1)}с`;
        } else if (evt.type === 'error') { result.value = 'Помилка: ' + evt.text; status.textContent = 'Помилка'; }
      }
    }
  } catch (e) {
    if (e.name !== 'AbortError') { log.textContent += 'Помилка запиту: ' + e + '\n'; status.textContent = 'Помилка'; }
  }
  setWorking(false); _controller = null; _requestId = null;
}

document.getElementById('input').addEventListener('keydown', e => {
  if (e.ctrlKey && e.key === 'Enter') doTranslate();
});

// ── File translation ──────────────────────────────────────────────────
let _fileController = null;
let _fileRequestId = null;
let _selectedFile = null;

function setFileWorking(on) {
  document.getElementById('file-btn').disabled = on;
  document.getElementById('file-stopBtn').style.display = on ? 'inline-block' : 'none';
  if (on) document.getElementById('download-link').style.display = 'none';
}

function fileSelected(input) {
  _selectedFile = input.files[0] || null;
  document.getElementById('file-name').textContent = _selectedFile ? _selectedFile.name : '';
}

function fileDragOver(e) { e.preventDefault(); document.getElementById('file-drop').classList.add('drag-over'); }
function fileDragLeave(e) { document.getElementById('file-drop').classList.remove('drag-over'); }
function fileDrop(e) {
  e.preventDefault();
  document.getElementById('file-drop').classList.remove('drag-over');
  const f = e.dataTransfer.files[0];
  if (f) { _selectedFile = f; document.getElementById('file-name').textContent = f.name; }
}

async function doFileStop() {
  if (_fileController) { _fileController.abort(); _fileController = null; }
  if (_fileRequestId) {
    await fetch('/stop/' + _fileRequestId, { method: 'POST' }).catch(() => {});
    _fileRequestId = null;
  }
  document.getElementById('file-status').textContent = 'Зупинено';
  setFileWorking(false);
  document.getElementById('progress-wrap').style.display = 'none';
}

async function doTranslateFile() {
  if (!_selectedFile) { document.getElementById('file-status').textContent = 'Спочатку оберіть файл'; return; }

  const fileLog = document.getElementById('file-log');
  const fileStatus = document.getElementById('file-status');
  const progressWrap = document.getElementById('progress-wrap');
  const progressBar = document.getElementById('progress-bar');

  setFileWorking(true);
  fileLog.style.display = 'block';
  fileLog.textContent = '';
  progressWrap.style.display = 'block';
  progressBar.style.width = '0%';
  fileStatus.textContent = 'Завантаження...';

  const startTime = Date.now();
  _fileController = new AbortController();

  const formData = new FormData();
  formData.append('file', _selectedFile, _selectedFile.name);
  formData.append('lang_from', document.getElementById('lang_from').value);
  formData.append('lang_to', document.getElementById('lang_to').value);

  try {
    const response = await fetch('/translate-file', {
      method: 'POST',
      signal: _fileController.signal,
      body: formData,
    });

    const reader = response.body.getReader();
    const decoder = new TextDecoder();
    let buffer = '';
    while (true) {
      const { done, value } = await reader.read();
      if (done) break;
      buffer += decoder.decode(value, { stream: true });
      const lines = buffer.split('\n');
      buffer = lines.pop();
      for (const line of lines) {
        if (!line.startsWith('data: ')) continue;
        let evt; try { evt = JSON.parse(line.slice(6)); } catch { continue; }
        if (evt.type === 'id') { _fileRequestId = evt.text; }
        else if (evt.type === 'log') { fileLog.textContent += evt.text + '\n'; fileLog.scrollTop = fileLog.scrollHeight; }
        else if (evt.type === 'progress') {
          fileStatus.textContent = evt.text;
          if (evt.pct !== undefined) progressBar.style.width = evt.pct + '%';
        } else if (evt.type === 'download') {
          const link = document.getElementById('download-link');
          link.href = evt.url;
          link.download = evt.filename;
          link.style.display = 'inline-block';
          progressBar.style.width = '100%';
          fileStatus.textContent = `Готово за ${((Date.now()-startTime)/1000).toFixed(1)}с`;
        } else if (evt.type === 'error') {
          fileStatus.textContent = 'Помилка: ' + evt.text;
          progressWrap.style.display = 'none';
        }
      }
    }
  } catch (e) {
    if (e.name !== 'AbortError') { fileLog.textContent += 'Помилка запиту: ' + e + '\n'; fileStatus.textContent = 'Помилка'; }
  }
  setFileWorking(false); _fileController = null; _fileRequestId = null;
}

// ── Stats ─────────────────────────────────────────────────────────────
async function loadStats() {
  try {
    const r = await fetch('/admin/stats?limit=100');
    const d = await r.json();
    const s = d.summary || {};
    const cards = [
      ['Всього', s.total ?? 0],
      ['Успішно', s.success ?? 0],
      ['Помилки', s.errors ?? 0],
      ['Зупинено', s.stopped ?? 0],
      ['Символів', s.total_chars ?? 0],
      ['Сторінок', s.total_pages ?? 0],
      ['Секунд', Math.round(s.total_seconds ?? 0)],
    ];
    document.getElementById('stats-summary').innerHTML = cards.map(([k, v]) =>
      `<div style="background:#f8fafc; border:1px solid #e2e8f0; border-radius:8px; padding:10px;">
         <div style="font-size:.7rem; color:#64748b; text-transform:uppercase;">${k}</div>
         <div style="font-size:1.2rem; font-weight:600; color:#1e293b; margin-top:2px;">${v}</div>
       </div>`).join('');

    const statusColor = {
      success: '#16a34a', error: '#dc2626', stopped: '#f59e0b',
    };
    const rows = (d.recent || []).map(r => {
      const t = (r.timestamp || '').replace('T', ' ').slice(5, 19);
      const langs = `${r.lang_from || ''}→${r.lang_to || ''}`;
      const color = statusColor[r.status] || '#64748b';
      return `<tr>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9;">${t}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9;">${r.ip || ''}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9;">${r.kind || ''}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9; max-width:200px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;">${r.filename || ''}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9;">${langs}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9; text-align:right;">${r.chars ?? ''}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9; text-align:right;">${r.pages ?? ''}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9; text-align:right;">${r.duration ?? ''}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9; color:${color}; font-weight:500;">${r.status}</td>
        <td style="padding:5px 8px; border-bottom:1px solid #f1f5f9; color:#dc2626; max-width:250px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;" title="${(r.error || '').replace(/"/g, '&quot;')}">${r.error || ''}</td>
      </tr>`;
    }).join('');
    document.getElementById('stats-tbody').innerHTML = rows || '<tr><td colspan="10" style="padding:20px; text-align:center; color:#94a3b8;">Немає даних</td></tr>';
  } catch (e) {
    document.getElementById('stats-tbody').innerHTML = `<tr><td colspan="10" style="padding:10px; color:#dc2626;">Помилка: ${e}</td></tr>`;
  }
}

document.getElementById('stats-details').addEventListener('toggle', e => {
  if (e.target.open) loadStats();
});

// ── Init selects ──────────────────────────────────────────────────────
async function initSelects() {
  const r = await fetch('/languages');
  const langs = await r.json();
  const from = document.getElementById('lang_from');
  const to   = document.getElementById('lang_to');
  from.appendChild(new Option('Автовизначення', 'auto'));
  langs.forEach(({ label, value }) => {
    from.appendChild(new Option(label, value));
    to.appendChild(new Option(label, value));
  });
  to.value = 'Ukrainian';
}

// ── Models & connection status ─────────────────────────────────────────
async function loadModels() {
  const dot  = document.getElementById('conn-dot');
  const txt  = document.getElementById('conn-text');
  const sel  = document.getElementById('model_select');
  try {
    const r    = await fetch('/models');
    const data = await r.json();
    sel.innerHTML = '';
    if (data.ok && data.models.length > 0) {
      dot.className = 'conn-dot ok';
      txt.textContent = 'Ollama підключена';
      data.models.forEach(m => {
        const opt = new Option(m, m);
        if (m === data.current) opt.selected = true;
        sel.appendChild(opt);
      });
      if (!data.models.includes(data.current) && data.current) {
        const opt = new Option(data.current + ' (не знайдено)', data.current);
        opt.selected = true;
        sel.insertBefore(opt, sel.firstChild);
      }
    } else {
      dot.className = 'conn-dot err';
      txt.textContent = 'Ollama недоступна';
      if (data.current) sel.appendChild(new Option(data.current, data.current));
    }
  } catch (e) {
    dot.className = 'conn-dot err';
    txt.textContent = 'Помилка з\'єднання';
  }
}

async function onModelChange() {
  const model = document.getElementById('model_select').value;
  await fetch('/config', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ model })
  });
}

setInterval(loadModels, 30000);

// ── Init ──────────────────────────────────────────────────────────────
initSelects();
loadModels();
loadSettings();
</script>
</body>
</html>"""


class TranslateRequest(BaseModel):
    text: str
    lang_from: str = "auto"
    lang_to: str


class TranslateHtmlRequest(BaseModel):
    html: str
    lang_from: str = "auto"
    lang_to: str


@app.get("/", response_class=HTMLResponse)
def index():
    return USER_HTML


@app.get("/admin", response_class=HTMLResponse)
def admin():
    return ADMIN_HTML


@app.get("/admin/stats")
def admin_stats(limit: int = 100):
    return JSONResponse({
        "summary": get_stats_summary(),
        "recent": get_recent_stats(limit),
    })



@app.get("/languages")
def get_languages():
    langs = sorted(LANG_MAP.keys(), key=lambda k: LANG_NAMES_UK[k])
    return JSONResponse([{"label": LANG_NAMES_UK[k], "value": k} for k in langs])


@app.get("/models")
def get_models():
    try:
        resp = req_lib.get(f"{_ollama_native_host()}/api/tags", timeout=5)
        if resp.status_code == 200:
            data = resp.json()
            models = [m["name"] for m in data.get("models", [])]
            return JSONResponse({"ok": True, "models": models, "current": CFG.get("model", "")})
        return JSONResponse({"ok": False, "models": [], "current": CFG.get("model", ""), "error": f"HTTP {resp.status_code}"})
    except Exception as e:
        return JSONResponse({"ok": False, "models": [], "current": CFG.get("model", ""), "error": str(e)})


@app.get("/config")
def get_config():
    return JSONResponse(CFG)


@app.get("/config/defaults")
def get_defaults():
    return JSONResponse(DEFAULTS)


@app.post("/config")
def post_config(data: dict):
    global CFG
    allowed = set(DEFAULTS.keys())
    for key in list(data.keys()):
        if key not in allowed:
            return JSONResponse({"ok": False, "error": f"Unknown key: {key}"}, status_code=400)
    CFG = {**CFG, **data}
    try:
        save_config(CFG)
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)
    return JSONResponse({"ok": True})


@app.post("/stop/{request_id}")
def stop_translation(request_id: str):
    event = _active.get(request_id)
    if event:
        event.set()
    return JSONResponse({"ok": True})


@app.post("/translate")
def translate(req: TranslateRequest, request: Request):
    global _ticket_issued
    request_id = str(uuid.uuid4())
    stop_event = threading.Event()
    _active[request_id] = stop_event
    client_ip = request.client.host if request.client else None

    start_evt = threading.Event()
    done_evt = threading.Event()
    with _ticket_lock:
        _ticket_issued += 1
        my_ticket = _ticket_issued
    _job_queue.put((start_evt, done_evt))

    def generate():
        def log_event(msg: str):
            return f"data: {json.dumps({'type': 'log', 'text': msg})}\n\n"

        def ts():
            return datetime.datetime.now().strftime("%H:%M:%S")

        started = time.time()
        final_status = "error"
        final_error = None

        try:
            yield f"data: {json.dumps({'type': 'id', 'text': request_id})}\n\n"

            # Length check up front
            max_chars = CFG.get("max_chars", 30000)
            if len(req.text) > max_chars:
                final_error = f"Text too long: {len(req.text)} > {max_chars}"
                msg = (f"Текст занадто довгий: {len(req.text)} символів "
                       f"(максимум {max_chars})")
                yield log_event(f"[{ts()}] ERROR: {msg}")
                yield f"data: {json.dumps({'type': 'error', 'text': msg})}\n\n"
                yield f"data: {json.dumps({'type': 'done'})}\n\n"
                return

            # Wait in queue until our turn
            while not start_evt.wait(timeout=2):
                if stop_event.is_set():
                    final_status = "stopped"
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    return
                with _ticket_lock:
                    ahead = my_ticket - _ticket_serving
                if ahead > 0:
                    yield f"data: {json.dumps({'type': 'queue', 'ahead': ahead})}\n\n"

            yield f"data: {json.dumps({'type': 'queue', 'ahead': 0})}\n\n"

            normalized = req.text.replace("\r\n", "\n")

            yield log_event(f"[{ts()}] URL:   {CFG['base_url']}")
            yield log_event(f"[{ts()}] Model: {CFG['model']}")
            yield log_event(f"[{ts()}] Lang:  {req.lang_from} → {req.lang_to}")
            yield log_event(f"[{ts()}] Text:  {len(req.text)} chars")

            try:
                # Single LLM call with whole text — preserves cross-paragraph
                # context (terminology, document domain). The model keeps the
                # paragraph breaks naturally in its output.
                translated = _translate_unit(
                    normalized, req.lang_from, req.lang_to, stop_event,
                )
                if translated is None:
                    final_status = "stopped"
                    yield log_event(f"[{ts()}] Stopped by user")
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    return

                final_status = "success"
                yield log_event(f"[{ts()}] Done — {len(translated)} chars")
                yield f"data: {json.dumps({'type': 'token', 'text': translated})}\n\n"
                yield f"data: {json.dumps({'type': 'result', 'text': translated})}\n\n"

            except req_lib.exceptions.Timeout:
                final_error = "Timeout"
                log.error(f"Translation timeout after {CFG['llm_timeout']}s")
                yield log_event(f"[{ts()}] ERROR: Timeout after {CFG['llm_timeout']}s")
                yield f"data: {json.dumps({'type': 'error', 'text': 'Сервер не встиг відповісти. Спробуйте пізніше.'})}\n\n"
            except Exception as e:
                final_error = str(e)
                log.error(f"Translation error: {e}\n{traceback.format_exc()}")
                yield log_event(f"[{ts()}] ERROR: {e}")
                yield log_event(traceback.format_exc())
                yield f"data: {json.dumps({'type': 'error', 'text': 'Помилка сервера. Деталі в адмін-логах.'})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        finally:
            done_evt.set()
            _active.pop(request_id, None)
            log_stat(
                ip=client_ip,
                kind="text",
                filename=None,
                file_ext=None,
                lang_from=req.lang_from,
                lang_to=req.lang_to,
                chars=len(req.text),
                pages=None,
                duration=round(time.time() - started, 2),
                status=final_status,
                error=final_error,
            )

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/translate-html")
def translate_html(req: TranslateHtmlRequest, request: Request):
    """Translate HTML preserving formatting tags — text nodes only."""
    global _ticket_issued
    request_id = str(uuid.uuid4())
    stop_event = threading.Event()
    _active[request_id] = stop_event
    client_ip = request.client.host if request.client else None

    start_evt = threading.Event()
    done_evt = threading.Event()
    with _ticket_lock:
        _ticket_issued += 1
        my_ticket = _ticket_issued
    _job_queue.put((start_evt, done_evt))

    def generate():
        def log_event(msg):
            return f"data: {json.dumps({'type': 'log', 'text': msg})}\n\n"
        def ts():
            return datetime.datetime.now().strftime("%H:%M:%S")

        started = time.time()
        final_status = "error"
        final_error = None

        try:
            yield f"data: {json.dumps({'type': 'id', 'text': request_id})}\n\n"

            max_chars = CFG.get("max_chars", 30000)
            if len(req.html) > max_chars:
                msg = f"Текст занадто довгий: {len(req.html)} символів (максимум {max_chars})"
                yield f"data: {json.dumps({'type': 'error', 'text': msg})}\n\n"
                yield f"data: {json.dumps({'type': 'done'})}\n\n"
                return

            while not start_evt.wait(timeout=2):
                if stop_event.is_set():
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    return
                with _ticket_lock:
                    ahead = my_ticket - _ticket_serving
                if ahead > 0:
                    yield f"data: {json.dumps({'type': 'queue', 'ahead': ahead})}\n\n"

            yield f"data: {json.dumps({'type': 'queue', 'ahead': 0})}\n\n"
            yield log_event(f"[{ts()}] HTML translation: {len(req.html)} bytes")

            try:
                import html2text as h2t

                # Convert HTML → Markdown (like TipTap does in OpenWebUI)
                converter = h2t.HTML2Text()
                converter.ignore_links = False
                converter.body_width = 0  # no line wrapping
                markdown_text = converter.handle(req.html).strip()

                if not markdown_text:
                    final_status = "error"
                    yield f"data: {json.dumps({'type': 'error', 'text': 'Порожній текст'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    return

                yield log_event(f"[{ts()}] Markdown: {len(markdown_text)} chars")

                # Stream translation of Markdown — model preserves ** ## - naturally
                translated_md = ""
                for token in _translate_unit_streaming(
                    markdown_text, req.lang_from, req.lang_to, stop_event
                ):
                    if stop_event.is_set():
                        final_status = "stopped"
                        yield f"data: {json.dumps({'type': 'done'})}\n\n"
                        return
                    translated_md += token
                    yield f"data: {json.dumps({'type': 'token', 'text': token})}\n\n"

                if not translated_md:
                    final_status = "error"
                    yield f"data: {json.dumps({'type': 'error', 'text': 'Ollama не повернув результат'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    return

                final_status = "success"
                yield log_event(f"[{ts()}] Done — {len(translated_md)} chars")
                # Send Markdown — frontend renders it via marked.js (like OpenWebUI)
                yield f"data: {json.dumps({'type': 'result', 'text': translated_md, 'format': 'markdown'})}\n\n"

            except req_lib.exceptions.Timeout:
                final_error = "Timeout"
                yield f"data: {json.dumps({'type': 'error', 'text': 'Сервер не встиг відповісти.'})}\n\n"
            except Exception as e:
                final_error = str(e)
                log.error(f"HTML translation error: {e}\n{traceback.format_exc()}")
                yield f"data: {json.dumps({'type': 'error', 'text': 'Помилка сервера.'})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        finally:
            done_evt.set()
            _active.pop(request_id, None)
            log_stat(
                ip=client_ip, kind="text", filename=None, file_ext=None,
                lang_from=req.lang_from, lang_to=req.lang_to,
                chars=len(req.html), pages=None,
                duration=round(time.time() - started, 2),
                status=final_status, error=final_error,
            )

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/translate-file")
async def translate_file_endpoint(
    request: Request,
    file: UploadFile = File(...),
    lang_from: str = Form("auto"),
    lang_to: str = Form("Ukrainian"),
):
    content = await file.read()
    filename = file.filename or "file.txt"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    base = filename.rsplit(".", 1)[0] if "." in filename else filename
    client_ip = request.client.host if request.client else None

    request_id = str(uuid.uuid4())
    stop_event = threading.Event()
    _active[request_id] = stop_event

    def generate():
        def ts():
            return datetime.datetime.now().strftime("%H:%M:%S")

        def log_event(msg: str):
            return f"data: {json.dumps({'type': 'log', 'text': f'[{ts()}] {msg}'})}\n\n"

        started = time.time()
        meta = {"chars": None, "pages": None}
        final_status = "error"
        final_error = None

        try:
            yield f"data: {json.dumps({'type': 'id', 'text': request_id})}\n\n"
            yield log_event(f"File: {filename} ({len(content)} bytes)")

            if ext == "docx":
                it = translate_docx_bytes(content, base, lang_from, lang_to, stop_event)
            elif ext == "pdf":
                it = translate_pdf_bytes(content, base, lang_from, lang_to, stop_event)
            else:
                it = translate_txt_bytes(content, base, lang_from, lang_to, stop_event)

            try:
                for event in it:
                    kind = event[0]
                    if kind == "log":
                        yield log_event(event[1])
                    elif kind == "meta":
                        meta.update(event[1])
                    elif kind == "progress":
                        yield f"data: {json.dumps({'type': 'progress', 'text': event[1], 'pct': event[2]})}\n\n"
                    elif kind == "error":
                        final_error = event[1]
                        log.warning(f"[{filename}] {event[1]}")
                        yield log_event(f"ERROR: {event[1]}")
                        yield f"data: {json.dumps({'type': 'error', 'text': event[1]})}\n\n"
                        return
                    elif kind == "stopped":
                        final_status = "stopped"
                        yield log_event("Зупинено користувачем")
                        yield f"data: {json.dumps({'type': 'done'})}\n\n"
                        return
                    elif kind == "done":
                        _, out_filename, mime, data = event
                        file_id = str(uuid.uuid4())
                        _results[file_id] = (out_filename, data, mime)
                        final_status = "success"
                        yield log_event(f"Готово — {len(data)} bytes")
                        yield f"data: {json.dumps({'type': 'download', 'url': f'/download/{file_id}', 'filename': out_filename})}\n\n"
                        yield f"data: {json.dumps({'type': 'done'})}\n\n"
                        return
            except req_lib.exceptions.Timeout:
                final_error = "Timeout від сервера перекладу"
                log.error(f"[{filename}] Timeout")
                yield log_event("ERROR: Timeout")
                yield f"data: {json.dumps({'type': 'error', 'text': 'Сервер не встиг відповісти. Спробуйте пізніше.'})}\n\n"
                return
            except Exception as e:
                final_error = str(e)
                log.error(f"[{filename}] {e}\n{traceback.format_exc()}")
                yield log_event(f"ERROR: {e}")
                yield log_event(traceback.format_exc())
                yield f"data: {json.dumps({'type': 'error', 'text': 'Помилка сервера. Деталі в адмін-логах.'})}\n\n"
                return

        finally:
            _active.pop(request_id, None)
            log_stat(
                ip=client_ip,
                kind="file",
                filename=filename,
                file_ext=ext,
                lang_from=lang_from,
                lang_to=lang_to,
                chars=meta.get("chars"),
                pages=meta.get("pages"),
                duration=round(time.time() - started, 2),
                status=final_status,
                error=final_error,
            )

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.get("/download/{file_id}")
def download_file(file_id: str):
    entry = _results.get(file_id)
    if not entry:
        return JSONResponse({"error": "not found"}, status_code=404)
    filename, data, mime = entry
    return Response(
        content=data,
        media_type=mime,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=HOST, port=PORT)
