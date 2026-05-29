import os
import sys
import datetime
import threading
import queue

# Bypass corporate proxy BEFORE any other imports
os.environ["NO_PROXY"] = "localhost,127.0.0.1"
os.environ["no_proxy"] = "localhost,127.0.0.1"
os.environ["HTTP_PROXY"] = ""
os.environ["HTTPS_PROXY"] = ""
os.environ["http_proxy"] = ""
os.environ["https_proxy"] = ""

import ollama
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
import customtkinter as ctk

VERSION = "3.6.0"
DEFAULT_OLLAMA_HOST = "http://127.0.0.1:11434"

POPULAR_LANGUAGES = [
    "Українська", "Англійська", "Німецька", "Іспанська", "Французька",
    "Польська", "Італійська", "Португальська", "Російська", "Чеська",
    "Румунська", "Болгарська", "Угорська", "Словацька", "Сербська", "Хорватська",
]

LANG_UK_TO_EN = {
    "Українська": "Ukrainian", "Англійська": "English", "Німецька": "German",
    "Іспанська": "Spanish", "Французька": "French", "Польська": "Polish",
    "Італійська": "Italian", "Португальська": "Portuguese", "Російська": "Russian",
    "Чеська": "Czech", "Румунська": "Romanian", "Болгарська": "Bulgarian",
    "Угорська": "Hungarian", "Словацька": "Slovak", "Сербська": "Serbian",
    "Хорватська": "Croatian",
}

ALL_LANGUAGES = [
    "Abkhazian", "Afar", "Afrikaans", "Akan", "Albanian", "Amharic", "Arabic",
    "Aragonese", "Armenian", "Assamese", "Azerbaijani", "Bambara", "Bashkir",
    "Basque", "Belarusian", "Bengali", "Bosnian", "Breton", "Bulgarian",
    "Burmese", "Catalan", "Central Khmer", "Chechen", "Chichewa", "Chinese",
    "Chuvash", "Cornish", "Corsican", "Croatian", "Czech", "Danish", "Divehi",
    "Dutch", "Dzongkha", "English", "Esperanto", "Estonian", "Ewe", "Faroese",
    "Filipino", "Finnish", "French", "Fulah", "Galician", "Georgian", "German",
    "Greek", "Guarani", "Gujarati", "Ganda", "Haitian", "Hausa", "Hebrew",
    "Hindi", "Hungarian", "Icelandic", "Ido", "Igbo", "Indonesian", "Interlingua",
    "Interlingue", "Inuktitut", "Inupiaq", "Irish", "Italian", "Japanese",
    "Javanese", "Kannada", "Kashmiri", "Kazakh", "Kikuyu", "Kinyarwanda",
    "Korean", "Kurdish", "Kyrgyz", "Kalaallisut", "Lao", "Latin", "Latvian",
    "Lingala", "Lithuanian", "Luba-Katanga", "Luxembourgish", "Macedonian",
    "Malagasy", "Malay", "Malayalam", "Maltese", "Manx", "Maori", "Marathi",
    "Mongolian", "Navajo", "Nepali", "North Ndebele", "Northern Sami",
    "Norwegian", "Norwegian Bokmål", "Norwegian Nynorsk", "Occitan", "Oriya",
    "Oromo", "Ossetian", "Pashto", "Persian", "Polish", "Portuguese", "Punjabi",
    "Quechua", "Romanian", "Romansh", "Rundi", "Russian", "Sango", "Sanskrit",
    "Sardinian", "Scottish Gaelic", "Serbian", "Shona", "Sichuan Yi", "Sindhi",
    "Sinhala", "Slovak", "Slovenian", "Somali", "South Ndebele", "Southern Sotho",
    "Spanish", "Sundanese", "Swahili", "Swati", "Swedish", "Tagalog", "Tamil",
    "Tatar", "Telugu", "Tibetan", "Tigrinya", "Tonga", "Tswana", "Tsonga",
    "Turkish", "Turkmen", "Uyghur", "Ukrainian", "Urdu", "Uzbek", "Venda",
    "Vietnamese", "Volapük", "Walloon", "Welsh", "Western Frisian", "Wolof",
    "Xhosa", "Yiddish", "Yoruba", "Zhuang", "Zulu",
]

if getattr(sys, 'frozen', False):
    _app_dir = os.path.dirname(sys.executable)
else:
    _app_dir = os.path.dirname(os.path.abspath(__file__))

_theme_path = os.path.join(_app_dir, "theme_custom.json")

_log_name = f"log_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
_log_path = os.path.join(_app_dir, _log_name)
_log_file = None
_logging_enabled = False


def log_file(msg):
    global _log_file
    if not _logging_enabled:
        return
    if _log_file is None:
        _log_file = open(_log_path, "w", encoding="utf-8")
    _log_file.write(msg + "\n")
    _log_file.flush()


# ── Translation logic ─────────────────────────────────────────────────────────

def translate_text(text, target_lang, model_name, host):
    if not text.strip() or len(text) < 2:
        return text
    prompt = (
        f"Translate the following text to {target_lang}. "
        f"Preserve the meaning and tone. Output ONLY the translation:\n\n{text}"
    )
    try:
        client = ollama.Client(host=host)
        response = client.generate(model=model_name, prompt=prompt)
        return response['response'].strip()
    except Exception as e:
        log_file(f"  [ERROR] {e}")
        return text


def run_translation(input_path, target_lang, model_name, host, msg_queue, stop_event):
    try:
        source_dir = os.path.dirname(input_path)
        filename = os.path.basename(input_path)
        output_path = os.path.join(source_dir, f"TR_{target_lang}_{filename}")

        msg_queue.put(("log", f"Відкриваю: {filename}"))
        log_file(f"[FILE] {input_path}")

        doc = Document(input_path)
        all_paras = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)

        translatable = [
            run
            for para in all_paras
            for run in para.runs
            if run.text.strip() and len(run.text.strip()) >= 2
        ]
        total = len(translatable)
        msg_queue.put(("log", f"Знайдено {total} текстових фрагментів для перекладу"))
        msg_queue.put(("total", total))

        for idx, run in enumerate(translatable, 1):
            if stop_event.is_set():
                msg_queue.put(("log", "⛔ Переклад зупинено користувачем"))
                msg_queue.put(("stopped", None))
                return

            original = run.text.strip()
            preview = original[:50] + ("..." if len(original) > 50 else "")
            msg_queue.put(("log", f"[{idx}/{total}] {preview}"))
            log_file(f"  [RUN {idx}] {repr(original[:80])}")

            translated = translate_text(original, target_lang, model_name, host)
            log_file(f"  [OUT]  {repr(translated[:80])}")

            if len(translated) > len(original) * 1.3 and run.font.size:
                try:
                    old_pt = run.font.size.pt
                    run.font.size = Pt(old_pt - 1)
                    log_file(f"  [FONT] {old_pt}pt → {old_pt - 1}pt")
                except Exception:
                    pass

            run.text = translated
            msg_queue.put(("progress", idx))

        doc.save(output_path)
        log_file(f"[FILE] Saved: {output_path}")
        msg_queue.put(("log", f"Збережено → TR_{target_lang}_{filename}"))
        msg_queue.put(("done", output_path))

    except Exception as e:
        log_file(f"[ERROR] {e}")
        msg_queue.put(("error", str(e)))


# ── All Languages Dialog ──────────────────────────────────────────────────────

class AllLanguagesDialog(ctk.CTkToplevel):
    def __init__(self, parent, on_select):
        super().__init__(parent)
        self.on_select = on_select
        self.title("Всі підтримувані мови")
        self.geometry("360x520")
        self.resizable(False, True)
        self.grab_set()
        self.lift()
        self.focus_force()

        ctk.CTkLabel(self, text="Пошук мови:", anchor="w").pack(fill="x", padx=12, pady=(12, 4))
        self._search_var = ctk.StringVar()
        self._search_var.trace_add("write", self._on_search)
        ctk.CTkEntry(self, textvariable=self._search_var,
                     placeholder_text="Введіть назву...").pack(fill="x", padx=12, pady=(0, 8))

        self._list_frame = ctk.CTkScrollableFrame(self, height=380)
        self._list_frame.pack(fill="both", expand=True, padx=12, pady=(0, 8))

        ctk.CTkButton(self, text="Скасувати", fg_color="gray40",
                      command=self.destroy).pack(pady=(0, 12))

        self._render_list(ALL_LANGUAGES)

    def _on_search(self, *_):
        q = self._search_var.get().lower()
        self._render_list([l for l in ALL_LANGUAGES if q in l.lower()])

    def _render_list(self, langs):
        for w in self._list_frame.winfo_children():
            w.destroy()
        for lang in langs:
            ctk.CTkButton(
                self._list_frame, text=lang, anchor="w",
                fg_color="transparent", hover_color=("gray75", "gray30"),
                text_color=("gray10", "gray90"),
                command=lambda l=lang: self._pick(l),
            ).pack(fill="x", pady=1)

    def _pick(self, lang):
        self.on_select(lang)
        self.destroy()


# ── Main App ──────────────────────────────────────────────────────────────────

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        self._theme = "dark"
        ctk.set_appearance_mode(self._theme)
        if os.path.isfile(_theme_path):
            ctk.set_default_color_theme(_theme_path)
        else:
            ctk.set_default_color_theme("blue")

        self.title(f"Перекладач DOCX  v{VERSION}")
        self.resizable(False, False)

        self._msg_queue = queue.Queue()
        self._total_runs = 0
        self._worker = None
        self._stop_event = threading.Event()
        self._selected_lang = POPULAR_LANGUAGES[0]
        self._log_visible = True
        self._server_visible = False

        self._build_ui()
        self._refresh_models()

    def _build_ui(self):
        pad = {"padx": 16, "pady": 5}

        # ── Header ──
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=16, pady=(14, 0))
        ctk.CTkLabel(
            header, text=f"Перекладач DOCX  v{VERSION}",
            font=ctk.CTkFont(size=18, weight="bold"),
        ).pack(side="left")
        theme_frame = ctk.CTkFrame(header, fg_color="transparent")
        theme_frame.pack(side="right")
        ctk.CTkLabel(theme_frame, text="☀", font=ctk.CTkFont(size=14)).pack(side="left", padx=(0, 2))
        self._theme_switch = ctk.CTkSwitch(theme_frame, text="", width=44, command=self._toggle_theme)
        self._theme_switch.pack(side="left")
        self._theme_switch.select()
        ctk.CTkLabel(theme_frame, text="🌙", font=ctk.CTkFont(size=14)).pack(side="left", padx=(2, 0))

        ctk.CTkLabel(
            self, text=f"Журнал: {_log_path}",
            font=ctk.CTkFont(size=11), text_color="gray",
        ).pack(pady=(2, 8))

        # ── File ──
        file_frame = ctk.CTkFrame(self, fg_color="transparent")
        file_frame.pack(fill="x", **pad)
        ctk.CTkLabel(file_frame, text="Файл:", width=80, anchor="w").pack(side="left")
        self._file_var = ctk.StringVar(value="Файл не вибрано")
        ctk.CTkEntry(file_frame, textvariable=self._file_var, state="readonly",
                     width=386).pack(side="left", padx=(0, 8))
        ctk.CTkButton(file_frame, text="Огляд", width=80, command=self._browse).pack(side="left")

        # ── Language ──
        lang_frame = ctk.CTkFrame(self, fg_color="transparent")
        lang_frame.pack(fill="x", **pad)
        ctk.CTkLabel(lang_frame, text="Мова:", width=80, anchor="w").pack(side="left")
        self._lang_menu = ctk.CTkOptionMenu(
            lang_frame, values=POPULAR_LANGUAGES, width=220,
            command=self._on_lang_select,
        )
        self._lang_menu.set(POPULAR_LANGUAGES[0])
        self._lang_menu.pack(side="left", padx=(0, 8))
        ctk.CTkButton(lang_frame, text="Всі мови...", width=110,
                      command=self._open_all_langs).pack(side="left")

        # ── Model ──
        model_frame = ctk.CTkFrame(self, fg_color="transparent")
        model_frame.pack(fill="x", **pad)
        ctk.CTkLabel(model_frame, text="Модель:", width=80, anchor="w").pack(side="left")
        self._model_var = ctk.StringVar(value="Завантаження...")
        self._model_menu = ctk.CTkOptionMenu(
            model_frame, variable=self._model_var, values=["Завантаження..."], width=386,
        )
        self._model_menu.pack(side="left", padx=(0, 8))
        ctk.CTkButton(model_frame, text="↺", width=40, command=self._refresh_models).pack(side="left")

        # ── Server settings (collapsible) ──
        srv_toggle_frame = ctk.CTkFrame(self, fg_color="transparent")
        srv_toggle_frame.pack(fill="x", padx=16, pady=(4, 0))
        self._srv_toggle_btn = ctk.CTkButton(
            srv_toggle_frame, text="▶  Налаштування сервера", anchor="w",
            fg_color="transparent", hover_color=("gray80", "gray25"),
            font=ctk.CTkFont(size=12), text_color=("gray40", "gray60"),
            command=self._toggle_server,
        )
        self._srv_toggle_btn.pack(side="left")

        self._server_frame = ctk.CTkFrame(self, fg_color=("gray90", "gray20"), corner_radius=8)

        srv_inner = ctk.CTkFrame(self._server_frame, fg_color="transparent")
        srv_inner.pack(fill="x", padx=12, pady=8)
        ctk.CTkLabel(srv_inner, text="Адреса Ollama:", width=110, anchor="w").pack(side="left")
        self._host_var = ctk.StringVar(value=DEFAULT_OLLAMA_HOST)
        ctk.CTkEntry(srv_inner, textvariable=self._host_var, width=320).pack(side="left", padx=(0, 8))
        ctk.CTkButton(srv_inner, text="Застосувати", width=110,
                      command=self._apply_server).pack(side="left")

        # ── Log toggle ──
        log_hdr = ctk.CTkFrame(self, fg_color="transparent")
        log_hdr.pack(fill="x", padx=16, pady=(8, 2))
        self._log_toggle_btn = ctk.CTkButton(
            log_hdr, text="▼  Журнал виконання", anchor="w",
            fg_color="transparent", hover_color=("gray80", "gray25"),
            font=ctk.CTkFont(size=12), text_color=("gray40", "gray60"),
            command=self._toggle_log,
        )
        self._log_toggle_btn.pack(side="left")
        self._log_enabled_var = ctk.BooleanVar(value=False)
        ctk.CTkCheckBox(
            log_hdr, text="Зберігати у файл", variable=self._log_enabled_var,
            font=ctk.CTkFont(size=12), command=self._toggle_logging,
        ).pack(side="right")

        self._log_box = ctk.CTkTextbox(
            self, height=180, font=ctk.CTkFont(family="Consolas", size=12),
            state="disabled", wrap="word",
        )
        self._log_box.pack(fill="x", padx=16)

        # ── Progress ──
        self._progress_label = ctk.CTkLabel(
            self, text="", text_color="gray", font=ctk.CTkFont(size=12),
        )
        self._progress_label.pack(pady=(8, 2))
        self._progressbar = ctk.CTkProgressBar(self, width=600)
        self._progressbar.set(0)
        self._progressbar.pack(padx=16, pady=(0, 10))

        # ── Buttons row ──
        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=(4, 16))
        self._start_btn = ctk.CTkButton(
            btn_frame, text="Розпочати переклад", height=42, width=280,
            font=ctk.CTkFont(size=14, weight="bold"),
            command=self._start,
        )
        self._start_btn.pack(side="left", padx=(0, 10))
        self._stop_btn = ctk.CTkButton(
            btn_frame, text="⛔ Стоп", height=42, width=120,
            font=ctk.CTkFont(size=13, weight="bold"),
            fg_color="#c0392b", hover_color="#96281b",
            state="disabled", command=self._stop,
        )
        self._stop_btn.pack(side="left")

    def _fit_window(self):
        self.update_idletasks()
        self.geometry(f"640x{self.winfo_reqheight()}")

    # ── Theme ──

    def _toggle_theme(self):
        self._theme = "dark" if self._theme_switch.get() else "light"
        ctk.set_appearance_mode(self._theme)

    # ── Server settings ──

    def _toggle_server(self):
        self._server_visible = not self._server_visible
        if self._server_visible:
            self._server_frame.pack(fill="x", padx=16, pady=(0, 4),
                                    before=self._log_toggle_btn.master)
            self._srv_toggle_btn.configure(text="▼  Налаштування сервера")
        else:
            self._server_frame.pack_forget()
            self._srv_toggle_btn.configure(text="▶  Налаштування сервера")
        self._fit_window()

    def _apply_server(self):
        host = self._host_var.get().strip()
        if not host:
            messagebox.showwarning("Порожня адреса", "Введіть адресу сервера Ollama.")
            return
        self._append_log(f"Адреса сервера змінена: {host}")
        self._refresh_models()

    def _get_host(self):
        return self._host_var.get().strip() or DEFAULT_OLLAMA_HOST

    # ── Log toggle ──

    def _toggle_log(self):
        self._log_visible = not self._log_visible
        if self._log_visible:
            self._log_box.pack(fill="x", padx=16,
                               before=self._progress_label)
            self._log_toggle_btn.configure(text="▼  Журнал виконання")
        else:
            self._log_box.pack_forget()
            self._log_toggle_btn.configure(text="▶  Журнал виконання")
        self._fit_window()

    def _toggle_logging(self):
        global _logging_enabled
        _logging_enabled = self._log_enabled_var.get()
        if _logging_enabled:
            self._append_log(f"Логування увімкнено → {_log_path}")
        else:
            self._append_log("Логування вимкнено")

    # ── Language ──

    def _on_lang_select(self, value):
        self._selected_lang = LANG_UK_TO_EN.get(value, value)

    def _open_all_langs(self):
        AllLanguagesDialog(self, self._set_custom_lang)

    def _set_custom_lang(self, lang):
        self._selected_lang = lang
        values = list(self._lang_menu.cget("values"))
        # Replace previous custom entry or add new one after popular list
        values = [v for v in values if v in POPULAR_LANGUAGES]
        values.append(lang)
        self._lang_menu.configure(values=values)
        self._lang_menu.set(lang)

    # ── Models ──

    def _refresh_models(self):
        self._model_menu.configure(state="disabled")
        self._model_var.set("Підключення...")

        def fetch():
            host = self._get_host()
            try:
                client = ollama.Client(host=host)
                resp = client.list()
                models = [m['model'] for m in resp.get('models', [])]
                if models:
                    self.after(0, lambda: self._set_models(models))
                else:
                    self.after(0, lambda: self._set_models_error("Моделі не знайдено"))
            except Exception as e:
                self.after(0, lambda: self._set_models_error(str(e)))

        threading.Thread(target=fetch, daemon=True).start()

    def _set_models(self, models):
        self._model_menu.configure(values=models, state="normal")
        self._model_var.set(models[0])
        self._append_log(f"Ollama підключено. Доступно моделей: {len(models)}")

    def _set_models_error(self, err):
        self._model_menu.configure(values=["— немає з'єднання —"], state="normal")
        self._model_var.set("— немає з'єднання —")
        self._append_log(f"[ПОМИЛКА] Ollama: {err}")

    # ── Log ──

    def _append_log(self, msg):
        self._log_box.configure(state="normal")
        self._log_box.insert("end", msg + "\n")
        self._log_box.see("end")
        self._log_box.configure(state="disabled")

    def _clear_log(self):
        self._log_box.configure(state="normal")
        self._log_box.delete("1.0", "end")
        self._log_box.configure(state="disabled")

    # ── File ──

    def _browse(self):
        path = filedialog.askopenfilename(
            title="Виберіть DOCX файл",
            filetypes=[("Word документи", "*.docx")],
        )
        if path:
            self._file_var.set(path)

    # ── Start / Stop ──

    def _start(self):
        input_path = self._file_var.get()
        if not input_path or input_path == "Файл не вибрано":
            messagebox.showwarning("Файл не вибрано", "Будь ласка, виберіть DOCX файл.")
            return
        if not os.path.isfile(input_path):
            messagebox.showerror("Файл не знайдено", f"Файл не існує:\n{input_path}")
            return
        model = self._model_var.get()
        if model in ("Завантаження...", "Підключення...", "— немає з'єднання —"):
            messagebox.showwarning("Модель не вибрана",
                                   "Дочекайтесь підключення до Ollama або оновіть список.")
            return

        self._clear_log()
        self._progressbar.set(0)
        self._progress_label.configure(text="")
        self._start_btn.configure(state="disabled", text="Перекладаю...")
        self._stop_btn.configure(state="normal")

        self._stop_event = threading.Event()
        self._msg_queue = queue.Queue()
        self._total_runs = 0
        self._worker = threading.Thread(
            target=run_translation,
            args=(input_path, self._selected_lang, model,
                  self._get_host(), self._msg_queue, self._stop_event),
            daemon=True,
        )
        self._worker.start()
        self.after(100, self._poll_queue)

    def _stop(self):
        self._stop_event.set()
        self._stop_btn.configure(state="disabled", text="Зупиняю...")

    def _poll_queue(self):
        try:
            while True:
                kind, data = self._msg_queue.get_nowait()
                if kind == "log":
                    self._append_log(data)
                elif kind == "total":
                    self._total_runs = data
                elif kind == "progress":
                    done = data
                    total = self._total_runs or 1
                    self._progressbar.set(done / total)
                    self._progress_label.configure(
                        text=f"{done} / {total}  ({int(done / total * 100)}%)"
                    )
                elif kind == "done":
                    self._on_done(data)
                    return
                elif kind == "stopped":
                    self._on_stopped()
                    return
                elif kind == "error":
                    self._on_error(data)
                    return
        except queue.Empty:
            pass
        self.after(100, self._poll_queue)

    def _on_done(self, output_path):
        self._progressbar.set(1)
        self._progress_label.configure(text="Готово!")
        self._append_log(f"\n✓ Переклад завершено!\n  {output_path}")
        self._start_btn.configure(state="normal", text="Розпочати переклад")
        self._stop_btn.configure(state="disabled", text="⛔ Стоп")
        messagebox.showinfo("Готово", f"Файл збережено:\n{output_path}")

    def _on_stopped(self):
        self._progress_label.configure(text="Зупинено")
        self._start_btn.configure(state="normal", text="Розпочати переклад")
        self._stop_btn.configure(state="disabled", text="⛔ Стоп")

    def _on_error(self, err):
        self._append_log(f"\n[ПОМИЛКА] {err}")
        self._start_btn.configure(state="normal", text="Розпочати переклад")
        self._stop_btn.configure(state="disabled", text="⛔ Стоп")
        messagebox.showerror("Помилка", err)


if __name__ == "__main__":
    app = App()
    app.mainloop()
    _log_file.close()
